import os
import sys
import logging
import uuid
import json
import re
import time
from collections import Counter, defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from functools import wraps
from datetime import datetime
from typing import Optional, List, Dict, Any
from pathlib import Path

import requests

from flask import Flask, jsonify, request, render_template_string, g, send_file
from pydantic import ValidationError

# 添加项目根目录到 Python 路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from processing import TextProcessor
from osint_cn.analysis import OSINTAnalyzer
from osint_cn.collection import CollectorFactory
from osint_cn.config import get_config
from osint_cn.models import (
    TextAnalysisRequest, SentimentRequest, KeywordRequest,
    CollectRequest, AnalyzeRequest, TaskCreateRequest,
    SentimentResponse, KeywordResponse, CollectResponse,
    HealthResponse, ErrorResponse, Platform
)
from osint_cn.security import (
    require_api_key, optional_api_key, rate_limit,
    setup_security_middleware, api_key_manager
)
from osint_cn.scheduler import (
    scheduler, TaskType, TaskConfig, TaskStatus
)
from osint_cn.realtime import (
    get_realtime_collector, get_streaming_collector,
    RealtimeCollector, CollectedItem
)
from osint_cn.intelligence import (
    get_intelligence_analyzer, IntelligenceAnalyzer,
    ThreatLevel, IntelligenceType, IntelligenceCategory
)
from osint_cn.relation import (
    get_knowledge_graph, get_social_network,
    KnowledgeGraph, SocialNetworkAnalyzer
)
from osint_cn.alert import (
    get_alert_manager, AlertManager, AlertLevel, AlertRule, RuleType, Alert, AlertStatus
)
from storage.service import (
    get_storage, CollectionRecord, CollectedItemRecord
)
from osint_cn.batch_collector import (
    get_scheduler, BatchCollectorScheduler, TaskStatus as BatchTaskStatus
)

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['JSON_AS_ASCII'] = False  # 支持中文 JSON

# 设置安全中间件
setup_security_middleware(app)

# 初始化组件
config = get_config()
text_processor = TextProcessor()
analyzer = OSINTAnalyzer(text_processor)

# 初始化核心模块
realtime_collector = get_realtime_collector()
intelligence_analyzer = get_intelligence_analyzer()
knowledge_graph = get_knowledge_graph()
social_network = get_social_network()
alert_manager = get_alert_manager()
storage_backend = get_storage()

def _env_flag(name: str, default: bool = False) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.lower() in ('1', 'true', 'yes', 'on')


ENABLE_SCHEDULER = _env_flag('ENABLE_SCHEDULER', True)
ENABLE_BACKGROUND_SERVICES = _env_flag('ENABLE_BACKGROUND_SERVICES', True)
ENFORCE_API_KEY = _env_flag('ENFORCE_API_KEY', False)


def _init_background_services() -> None:
    if ENABLE_SCHEDULER:
        try:
            scheduler.start()
            logger.info('Scheduler started')
        except Exception as exc:
            logger.exception(f'Failed to start scheduler: {exc}')
    else:
        logger.info('Scheduler startup disabled by ENABLE_SCHEDULER')

    if ENABLE_BACKGROUND_SERVICES:
        try:
            realtime_collector.start()
            logger.info('Realtime collector started')
        except Exception as exc:
            logger.exception(f'Failed to start realtime collector: {exc}')

        try:
            alert_manager.start()
            logger.info('Alert manager started')
        except Exception as exc:
            logger.exception(f'Failed to start alert manager: {exc}')
    else:
        logger.info('Background services startup disabled by ENABLE_BACKGROUND_SERVICES')


_init_background_services()

# 数据存储（内存，生产环境应使用数据库）
collected_data_store = {}
analysis_results_store = {}

# 仪表板一体化分析结果（内存）
dashboard_pipeline_store = {}
monitor_profiles_store = {}
monitor_groups_store = {}


@dataclass
class MonitorGroup:
    """监控对象分组。"""
    group_id: str
    name: str
    description: str = ''
    color: str = '#2e8cff'
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    updated_at: str = field(default_factory=lambda: datetime.now().isoformat())

    def to_dict(self) -> Dict[str, Any]:
        return {
            'group_id': self.group_id,
            'name': self.name,
            'description': self.description,
            'color': self.color,
            'created_at': self.created_at,
            'updated_at': self.updated_at
        }


@dataclass
class MonitorProfile:
    """监控对象配置。"""
    monitor_id: str
    name: str
    keywords: List[str]
    platforms: List[str]
    group_id: Optional[str] = None
    tags: List[str] = field(default_factory=list)
    interval_seconds: int = 1800
    max_items: int = 60
    thresholds: Dict[str, Any] = field(default_factory=dict)
    report_formats: List[str] = field(default_factory=lambda: ['docx', 'pdf'])
    enabled: bool = True
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    updated_at: str = field(default_factory=lambda: datetime.now().isoformat())
    last_run_at: Optional[str] = None
    last_status: str = 'idle'
    last_pipeline_ids: List[str] = field(default_factory=list)
    scheduler_job_id: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            'monitor_id': self.monitor_id,
            'name': self.name,
            'keywords': self.keywords,
            'platforms': self.platforms,
            'group_id': self.group_id,
            'tags': self.tags,
            'interval_seconds': self.interval_seconds,
            'max_items': self.max_items,
            'thresholds': self.thresholds,
            'report_formats': self.report_formats,
            'enabled': self.enabled,
            'created_at': self.created_at,
            'updated_at': self.updated_at,
            'last_run_at': self.last_run_at,
            'last_status': self.last_status,
            'last_pipeline_ids': self.last_pipeline_ids
        }


def _ensure_export_dir() -> Path:
    export_dir = Path(os.getcwd()) / 'logs' / 'exports'
    export_dir.mkdir(parents=True, exist_ok=True)
    return export_dir


def _normalize_keywords(keywords: Any) -> List[str]:
    if isinstance(keywords, str):
        parts = re.split(r'[,，\n]+', keywords)
    else:
        parts = keywords or []
    normalized = []
    for item in parts:
        keyword = str(item or '').strip()
        if keyword and keyword not in normalized:
            normalized.append(keyword)
    return normalized


def _normalize_tags(tags: Any) -> List[str]:
    if isinstance(tags, str):
        parts = re.split(r'[,，\n]+', tags)
    else:
        parts = tags or []
    normalized = []
    for item in parts:
        tag = str(item or '').strip()
        if tag and tag not in normalized:
            normalized.append(tag)
    return normalized


def _normalize_platforms(platforms: Optional[List[str]]) -> List[str]:
    """标准化平台列表并过滤非法平台。"""
    available = set(CollectorFactory.available_platforms())
    if not platforms:
        return ['weibo', 'zhihu', 'baidu']

    normalized = []
    for platform in platforms:
        if not platform:
            continue
        key = str(platform).strip().lower()
        if key in available and key not in normalized:
            normalized.append(key)

    if normalized:
        return normalized
    return ['weibo', 'zhihu', 'baidu']


def _generate_wordcloud_from_items(items: List[Dict[str, Any]], top_k: int = 80) -> List[Dict[str, Any]]:
    """根据采集内容生成词云数据。"""
    text = ' '.join((item.get('content') or '').strip() for item in items)
    if not text:
        return []

    frequencies = text_processor.word_frequency(text, top_k=top_k)
    return [{'name': word, 'value': freq} for word, freq in frequencies if word]


def _build_report_html(
    keyword: str,
    platforms: List[str],
    total_items: int,
    analysis: Dict[str, Any],
    wordcloud: List[Dict[str, Any]],
    ai_report: Optional[Dict[str, Any]] = None
) -> str:
    """生成专业的HTML报告，格式与参考模板相同。"""
    # 提取数据
    sentiment = analysis.get('sentiment', {}).get('data', {}).get('statistics', {})
    risk = analysis.get('risk', {}).get('data', {})
    trend = analysis.get('trend', {}).get('data', {})

    positive = sentiment.get('positive_count', 0)
    negative = sentiment.get('negative_count', 0)
    neutral = sentiment.get('neutral_count', 0)
    risk_level = risk.get('risk_level', 'low')
    risk_score = risk.get('risk_score', 0)
    peak_time = trend.get('peak_time') or '未知'
    peak_count = trend.get('peak_count', 0)

    hot_words = [item.get('name') for item in wordcloud[:10] if item.get('name')]
    risk_recommendations = risk.get('recommendations', [])

    # 处理 ai_report
    ai_report = ai_report or {}
    if isinstance(ai_report, str):
        ai_summary = ''
    else:
        ai_summary = ai_report.get('executive_summary', '') if isinstance(ai_report, dict) else ''

    # 生成Sentiment数据用于Chart.js
    sentiment_data = {
        'positive': positive,
        'neutral': neutral,
        'negative': negative
    }

    # 计算比例
    total = positive + negative + neutral + 1
    positive_ratio = (positive / total) * 100
    neutral_ratio = (neutral / total) * 100
    negative_ratio = (negative / total) * 100

    # 生成KPI数据
    kpi_cards = [
        {'label': '正向评论', 'value': positive, 'change': f"+{positive}", 'type': 'positive'},
        {'label': '中立评论', 'value': neutral, 'change': f"{neutral}", 'type': 'neutral'},
        {'label': '负向评论', 'value': negative, 'change': f"{negative}", 'type': 'negative'},
        {'label': '总体样本', 'value': total_items, 'change': f"同比采集", 'type': 'primary'},
        {'label': '风险等级', 'value': risk_level.upper(), 'change': f"分值 {risk_score}", 'type': 'risk'},
        {'label': '覆盖平台', 'value': len(platforms), 'change': f"{', '.join(platforms[:2])}", 'type': 'primary'},
    ]

    # 生成热词列表
    hot_words_html = ', '.join(hot_words) if hot_words else '暂无明显热词'

    # 定义平台描述（在f-string外部）
    platform_desc_dict = {
        'weibo': '国内最大的微博社交平台，舆论传播的主阵地',
        'douyin': '短视频平台，情绪传播最速渠道',
        'kuaishou': '快手短视频平台，下沉市场重要渠道',
        'zhihu': '知识问答社区，深度讨论的重要阵地',
        'baidu': '百度搜索与百家号，信息整合平台',
        'wechat': '微信公众号与朋友圈，熟人网络传播',
        'xiaohongshu': '小红书社区，Z世代用户集中地',
        'bilibili': 'B站视频社区，弹幕文化与亚文化阵地',
        'tieba': '百度贴吧，垂直社区讨论平台',
        'toutiao': '今日头条，信息流推荐平台'
    }

    # 生成平台HTML
    platform_rows_html = ''
    for platform in platforms:
        desc = platform_desc_dict.get(platform, '社交媒体平台')
        platform_rows_html += f"""                            <tr>
                                <td><strong>{platform}</strong></td>
                                <td>{desc}</td>
                            </tr>
"""

    # 生成AI洞察块
    ai_insight_html = ''
    if ai_summary:
        ai_insight_html = f'<div class="highlight-box"><h4>AI智能分析洞察</h4><p>{ai_summary}</p></div>'

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{keyword} - 舆情与客诉分析报告</title>
    <script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.0/dist/chart.umd.js"></script>
    <style>
        :root {{
            --primary-color: #2c3e50;
            --accent-color: #3498db;
            --background-color: #ffffff;
            --text-color: #2c3e50;
            --secondary-color: #7f8c8d;
            --card-background: #f8f9fa;
            --border-color: #e1e8ed;
            --success-color: #27ae60;
            --warning-color: #f39c12;
            --danger-color: #c0392b;
            --table-header-bg: #ecf0f1;
            --code-bg: #f5f5f5;
            --quote-border: #bdc3c7;
        }}

        [data-theme="dark"] {{
            --primary-color: #ecf0f1;
            --accent-color: #3498db;
            --background-color: #1a1a1a;
            --text-color: #ecf0f1;
            --secondary-color: #bdc3c7;
            --card-background: #2a2a2a;
            --border-color: #444444;
            --table-header-bg: #333333;
            --code-bg: #2a2a2a;
            --quote-border: #555555;
        }}

        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        html {{
            scroll-behavior: smooth;
        }}

        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Helvetica Neue', 'Microsoft YaHei', sans-serif;
            background-color: var(--background-color);
            color: var(--text-color);
            line-height: 1.6;
            transition: background-color 0.3s, color 0.3s;
        }}

        .container {{
            max-width: 1200px;
            margin: 0 auto;
            padding: 0 20px;
        }}

        header {{
            background: linear-gradient(135deg, var(--primary-color) 0%, var(--accent-color) 100%);
            color: white;
            padding: 40px 0;
            margin-bottom: 40px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        }}

        .controls {{
            display: flex;
            justify-content: flex-end;
            gap: 10px;
            margin-bottom: 20px;
        }}

        .control-btn {{
            padding: 8px 16px;
            background-color: rgba(255,255,255,0.2);
            color: white;
            border: 1px solid rgba(255,255,255,0.3);
            border-radius: 4px;
            cursor: pointer;
            font-size: 0.9em;
            transition: all 0.3s;
        }}

        .control-btn:hover {{
            background-color: rgba(255,255,255,0.3);
            border-color: rgba(255,255,255,0.5);
        }}

        .report-title {{
            font-size: 2.5em;
            font-weight: 700;
            margin-bottom: 10px;
        }}

        .report-subtitle {{
            font-size: 1.1em;
            opacity: 0.9;
        }}

        main {{
            padding: 20px 0;
        }}

        .toc {{
            background-color: var(--card-background);
            border: 1px solid var(--border-color);
            padding: 30px;
            border-radius: 8px;
            margin-bottom: 40px;
        }}

        .toc h2 {{
            font-size: 1.3em;
            margin-bottom: 20px;
            color: var(--primary-color);
        }}

        .toc ul {{
            list-style: none;
        }}

        .toc li {{
            margin: 10px 0;
        }}

        .toc > ul > li {{
            margin-left: 0;
        }}

        .toc > ul > li > ul > li {{
            margin-left: 20px;
            font-size: 0.95em;
        }}

        .toc a {{
            color: var(--accent-color);
            text-decoration: none;
            transition: color 0.3s;
        }}

        .toc a:hover {{
            color: var(--primary-color);
            text-decoration: underline;
        }}

        .content-section {{
            background-color: var(--card-background);
            border: 1px solid var(--border-color);
            padding: 30px;
            margin-bottom: 30px;
            border-radius: 8px;
        }}

        h1 {{
            font-size: 2em;
            margin-bottom: 20px;
            color: var(--primary-color);
            border-bottom: 3px solid var(--accent-color);
            padding-bottom: 10px;
        }}

        h2 {{
            font-size: 1.6em;
            margin-top: 30px;
            margin-bottom: 15px;
            color: var(--primary-color);
        }}

        h3 {{
            font-size: 1.2em;
            margin-top: 20px;
            margin-bottom: 10px;
            color: var(--accent-color);
        }}

        h4 {{
            font-size: 1.05em;
            margin-top: 15px;
            margin-bottom: 10px;
            color: var(--secondary-color);
            font-weight: 600;
        }}

        p {{
            margin-bottom: 15px;
            line-height: 1.8;
        }}

        ul, ol {{
            margin-left: 30px;
            margin-bottom: 15px;
        }}

        li {{
            margin-bottom: 8px;
        }}

        .chart-container {{
            position: relative;
            margin: 30px auto;
            height: 400px;
            width: 100%;
            max-width: 800px;
        }}

        table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 2em;
            font-size: 0.95em;
        }}

        th, td {{
            padding: 12px 15px;
            border: 1px solid var(--border-color);
            text-align: left;
        }}

        th {{
            background-color: var(--table-header-bg);
            font-weight: 600;
        }}

        tr:nth-child(even) {{
            background-color: var(--code-bg);
        }}

        blockquote {{
            border-left: 5px solid var(--quote-border);
            padding: 15px 20px;
            margin: 20px 0;
            background-color: var(--code-bg);
            font-style: italic;
            color: var(--secondary-color);
        }}

        .highlight-box {{
            background-color: rgba(52, 152, 219, 0.1);
            border: 1px solid rgba(52, 152, 219, 0.3);
            padding: 20px;
            margin: 20px 0;
            border-radius: 5px;
        }}

        .highlight-box h4 {{
            color: var(--accent-color);
            margin-top: 0;
        }}

        .kpi-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
            margin-top: 20px;
        }}

        .kpi-card {{
            background-color: var(--background-color);
            padding: 20px;
            border-radius: 8px;
            text-align: center;
            border: 1px solid var(--border-color);
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }}

        .kpi-value {{
            font-size: 2.2em;
            font-weight: 700;
            color: var(--accent-color);
            margin-bottom: 5px;
        }}

        .kpi-label {{
            font-size: 0.95em;
            color: var(--secondary-color);
            margin-bottom: 8px;
        }}

        .kpi-change {{
            font-size: 0.9em;
            font-weight: 500;
        }}

        .kpi-change.positive {{ color: var(--success-color); }}
        .kpi-change.negative {{ color: var(--danger-color); }}
        .kpi-change.neutral {{ color: var(--secondary-color); }}

        .engine-perspective {{
            border: 1px dashed var(--border-color);
            padding: 15px;
            margin: 20px 0;
            border-radius: 5px;
            background-color: rgba(52, 152, 219, 0.05);
        }}

        .engine-name {{
            font-weight: bold;
            color: var(--accent-color);
            display: block;
            margin-bottom: 10px;
            font-size: 0.9em;
        }}

        @media (max-width: 768px) {{
            .report-title {{ font-size: 2em; }}
            .report-subtitle {{ font-size: 1em; }}
            .controls {{ position: static; margin-top: 20px; justify-content: center; }}
            .content-section, .toc {{ padding: 20px; }}
            h1 {{ font-size: 1.8em; }}
            h2 {{ font-size: 1.5em; }}
            h3 {{ font-size: 1.2em; }}
            .chart-container {{ max-width: 100%; }}
            .kpi-grid {{ grid-template-columns: 1fr; }}
        }}

        @media print {{
            body {{
                --background-color: #ffffff;
                --text-color: #000000;
                --card-background: #ffffff;
                --border-color: #cccccc;
                --primary-color: #000000;
                --secondary-color: #333333;
                --accent-color: #000000;
                box-shadow: none;
            }}
            header, .controls, .toc {{
                display: none;
            }}
            .content-section {{
                box-shadow: none;
                border: 1px solid #ccc;
                page-break-inside: avoid;
            }}
            .chart-container {{
                page-break-inside: avoid;
            }}
            table, blockquote {{
                page-break-inside: avoid;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <header>
            <div class="controls">
                <button id="theme-toggle" class="control-btn" onclick="toggleTheme()">暗色模式</button>
                <button id="print-btn" class="control-btn" onclick="window.print()">打印/导出PDF</button>
            </div>
            <h1 class="report-title">{keyword} 舆情与客诉分析报告</h1>
            <p class="report-subtitle">基于多渠道数据采集与深度AI分析的综合评估</p>
        </header>

        <main>
            <nav class="toc" id="table-of-contents">
                <h2>报告目录</h2>
                <ul>
                    <li><a href="#section-1">1. 摘要与核心指标</a>
                        <ul>
                            <li><a href="#section-1-1">1.1 监测概览</a></li>
                            <li><a href="#section-1-2">1.2 关键指标表现</a></li>
                        </ul>
                    </li>
                    <li><a href="#section-2">2. 情感与分布分析</a>
                        <ul>
                            <li><a href="#section-2-1">2.1 情感结构</a></li>
                            <li><a href="#section-2-2">2.2 渠道分布</a></li>
                            <li><a href="#section-2-3">2.3 热点词汇</a></li>
                        </ul>
                    </li>
                    <li><a href="#section-3">3. 风险评估与建议</a>
                        <ul>
                            <li><a href="#section-3-1">3.1 风险研判</a></li>
                            <li><a href="#section-3-2">3.2 处置建议</a></li>
                        </ul>
                    </li>
                </ul>
            </nav>

            <section id="section-1" class="content-section">
                <h1>1. 摘要与核心指标</h1>

                <h2 id="section-1-1">1.1 监测概览</h2>
                <p>本报告对关键词"{keyword}"进行了全面监测分析，覆盖{len(platforms)}个主流社交媒体平台：{', '.join(platforms)}。采集有效数据{total_items}条，运用多维度分析模型，为您提供专业的舆情与客诉洞察。</p>
                {ai_insight_html}

                <h2 id="section-1-2">1.2 关键指标表现</h2>
                <div class="kpi-grid">
"""
    for card in kpi_cards:
        change_class = card.get('type', 'primary')
        html += f"""                    <div class="kpi-card">
                        <div class="kpi-value">{card['value']}</div>
                        <div class="kpi-label">{card['label']}</div>
                        <div class="kpi-change {change_class}">{card.get('change', '')}</div>
                    </div>
"""

    html += f"""                </div>

            </section>

            <section id="section-2" class="content-section">
                <h1>2. 情感与分布分析</h1>

                <h2 id="section-2-1">2.1 情感结构</h2>
                <div class="chart-container">
                    <canvas id="sentimentChart"></canvas>
                </div>
                <p style="text-align: center; margin-top: 15px; font-size: 0.9em; color: var(--secondary-color);">
                    本次监测中，正向声量占{positive_ratio:.1f}%，
                    中立声量占{neutral_ratio:.1f}%，
                    负向声量占{negative_ratio:.1f}%。
                </p>

                <h2 id="section-2-2">2.2 渠道分布</h2>
                <div class="highlight-box">
                    <h4>监测平台列表</h4>
                    <p>本次分析涵盖了以下主流社交媒体和资讯平台：</p>
                    <table>
                        <thead>
                            <tr>
                                <th>平台名称</th>
                                <th>说明</th>
                            </tr>
                        </thead>
                        <tbody>
{platform_rows_html}                        </tbody>
                    </table>
                </div>

                <h2 id="section-2-3">2.3 热点词汇</h2>
                <p>根据词频统计与热度分析，监测周期内最受关注的词汇包括：</p>
                <blockquote>
                    {hot_words_html}
                </blockquote>
                <p style="margin-top: 10px; font-size: 0.9em; color: var(--secondary-color);">
                    这些热词反映了公众最关注的焦点与舆论导向。建议重点关注负面热词并制定针对性的回应策略。
                </p>

            </section>

            <section id="section-3" class="content-section">
                <h1>3. 风险评估与建议</h1>

                <h2 id="section-3-1">3.1 风险研判</h2>
                <div class="highlight-box">
                    <h4>风险等级评估</h4>
                    <p><strong>当前风险等级：<span style="color: var(--accent-color); font-size: 1.1em;">{risk_level.upper()}</span></strong></p>
                    <p>综合风险分数：<strong>{risk_score}</strong></p>
                    <p>传播峰值时段：<strong>{peak_time}</strong>（{peak_count}条）</p>
                </div>

                <h2 id="section-3-2">3.2 处置建议</h2>
                <div class="engine-perspective">
                    <span class="engine-name">🔍 数据驱动建议</span>
                    <ol>
"""
    
    if risk_recommendations:
        for rec in risk_recommendations[:5]:
            html += f"                        <li>{rec}</li>\n"
    else:
        html += """                        <li>建立 7x24 关键词监测与分级告警规则。</li>
                        <li>对高传播负面内容在 2 小时内完成首轮响应。</li>
                        <li>将客服、运营、品牌部门纳入同一事件协同群。</li>
                        <li>定期发布官方声明，主动设置议题。</li>
                        <li>建立客诉-舆情关联闭环追踪机制。</li>
"""

    html += """                    </ol>
                </div>

            </section>

        </main>
    </div>

    <script>
        // 切换深色/浅色主题
        function toggleTheme() {{
            const html = document.documentElement;
            const btn = document.getElementById('theme-toggle');
            const currentTheme = html.getAttribute('data-theme');
            const newTheme = currentTheme === 'dark' ? 'light' : 'dark';
            html.setAttribute('data-theme', newTheme);
            localStorage.setItem('reportTheme', newTheme);
            btn.textContent = newTheme === 'dark' ? '浅色模式' : '暗色模式';
        }}

        // 初始化主题
        (function initTheme() {{
            const saved = localStorage.getItem('reportTheme') || 'light';
            document.documentElement.setAttribute('data-theme', saved);
            document.getElementById('theme-toggle').textContent = saved === 'dark' ? '浅色模式' : '暗色模式';
        }})();

        // Chart.js：情感分布
        const sentimentCtx = document.getElementById('sentimentChart')?.getContext('2d');
        if (sentimentCtx) {{
            new Chart(sentimentCtx, {{
                type: 'doughnut',
                data: {{
                    labels: ['正向 ({positive}条)', '中立 ({neutral}条)', '负向 ({negative}条)'],
                    datasets: [{{
                        data: [{positive}, {neutral}, {negative}],
                        backgroundColor: ['#27ae60', '#95a5a6', '#e74c3c'],
                        borderColor: 'var(--background-color)',
                        borderWidth: 2,
                        borderRadius: [8, 8, 8]
                    }}]
                }},
                options: {{
                    responsive: true,
                    maintainAspectRatio: false,
                    plugins: {{
                        legend: {{
                            position: 'bottom',
                            labels: {{
                                color: 'var(--text-color)',
                                font: {{ size: 12 }}
                            }}
                        }},
                        tooltip: {{
                            borderColor: 'var(--border-color)',
                            backgroundColor: 'var(--card-background)',
                            titleColor: 'var(--text-color)',
                            bodyColor: 'var(--text-color)',
                            borderWidth: 1
                        }}
                    }}
                }}
            }});
        }}
    </script>
</body>
</html>
"""
    return html


def _build_report_text(
    keyword: str,
    platforms: List[str],
    total_items: int,
    analysis: Dict[str, Any],
    wordcloud: List[Dict[str, Any]],
    ai_report: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """组装舆情/客诉报告主体。"""
    sentiment = analysis.get('sentiment', {}).get('data', {}).get('statistics', {})
    risk = analysis.get('risk', {}).get('data', {})
    trend = analysis.get('trend', {}).get('data', {})

    positive = sentiment.get('positive_count', 0)
    negative = sentiment.get('negative_count', 0)
    neutral = sentiment.get('neutral_count', 0)
    risk_level = risk.get('risk_level', 'low')
    risk_score = risk.get('risk_score', 0)
    peak_time = trend.get('peak_time') or '未知'
    peak_count = trend.get('peak_count', 0)

    hot_words = [item.get('name') for item in wordcloud[:10] if item.get('name')]
    risk_recommendations = risk.get('recommendations', [])

    findings = [
        f"本次监测关键词“{keyword}”共采集 {total_items} 条多渠道内容，覆盖平台：{', '.join(platforms)}。",
        f"情绪结构：正面 {positive}、中立 {neutral}、负面 {negative}。",
        f"风险等级：{risk_level}（风险分 {risk_score}），传播峰值时段：{peak_time}（{peak_count} 条）。",
        f"高频关注词：{('、'.join(hot_words) if hot_words else '暂无明显热词')}。"
    ]

    recommendations = list(risk_recommendations)
    if not recommendations:
        recommendations = [
            '建议建立 7x24 关键词监测与分级告警规则。',
            '建议对高传播负面内容在 2 小时内完成首轮响应。',
            '建议将客诉工单与舆情事件进行关联闭环追踪。'
        ]

    ai_report = ai_report or {}
    # 处理 ai_report 可能是字符串的情况（在零数据场景下）
    if isinstance(ai_report, str):
        ai_actions = []
    else:
        ai_actions = ai_report.get('action_recommendations', []) if isinstance(ai_report, dict) else []
    if ai_actions:
        for item in ai_actions[:3]:
            if item not in recommendations:
                recommendations.append(item)

    summary = (
        f"围绕“{keyword}”的全网舆情总体风险为 {risk_level}，"
        f"负面声量为 {negative} 条，建议聚焦高频客诉主题并快速响应。"
    )

    return {
        'title': f"{keyword} 舆情与客诉分析报告",
        'generated_at': datetime.now().isoformat(),
        'summary': summary,
        'findings': findings,
        'recommendations': recommendations,
        'ai_enhanced': bool(ai_report.get('enhanced')) if isinstance(ai_report, dict) else False,
        'ai_insight': {
            'executive_summary': ai_report.get('executive_summary', '') if isinstance(ai_report, dict) else '',
            'risk_judgment': ai_report.get('risk_judgment', '') if isinstance(ai_report, dict) else '',
            'action_recommendations': ai_actions,
            'pr_talking_points': ai_report.get('pr_talking_points', []) if isinstance(ai_report, dict) else [],
            'source': ai_report.get('source', 'rule_based') if isinstance(ai_report, dict) else 'diagnostics'
        }
    }


def _build_rule_based_ai_report(report_context: Dict[str, Any]) -> Dict[str, Any]:
    """外部模型不可用时的规则兜底报告。"""
    keyword = report_context.get('keyword', '目标对象')
    risk = report_context.get('risk', {}).get('data', {})
    risk_level = risk.get('risk_level', 'low')
    risk_score = risk.get('risk_score', 0)
    platforms = report_context.get('platforms', [])
    wordcloud = report_context.get('wordcloud', [])
    hot_words = [item.get('name') for item in wordcloud[:5] if item.get('name')]

    risk_map = {
        'critical': '事件已接近危机传播态势，应立即启动专项响应。',
        'high': '负面扩散速度较快，需要统一口径并同步业务部门。',
        'medium': '舆情已形成可见讨论面，建议尽快完成解释与安抚。',
        'low': '当前风险可控，以持续监测和定点回应为主。'
    }

    return {
        'enhanced': False,
        'source': 'rule_based',
        'executive_summary': (
            f"围绕“{keyword}”的讨论已覆盖 {len(platforms)} 个渠道，"
            f"当前风险等级为 {risk_level}，建议以监测和快速响应并行推进。"
        ),
        'risk_judgment': risk_map.get(risk_level, '当前风险可控，建议继续监测。') + f" 综合风险分为 {risk_score}。",
        'action_recommendations': [
            f"优先核查“{keyword}”相关客诉与负面反馈来源。",
            '2小时内完成首轮回应，24小时内输出处理进展。',
            '将客服、运营、品牌部门纳入同一事件协同群。'
        ],
        'pr_talking_points': [
            f"我们已关注到与“{keyword}”相关的反馈，并已启动专项核查。",
            '对于已确认的问题将尽快处理，并同步公开进展。',
            f"针对高频关注点{('：' + '、'.join(hot_words)) if hot_words else ''}，将提供更明确说明。"
        ]
    }


def _extract_json_object(content: str) -> Optional[Dict[str, Any]]:
    """从模型文本中提取 JSON 对象。"""
    if not content:
        return None

    stripped = content.strip()
    try:
        return json.loads(stripped)
    except Exception:
        pass

    match = re.search(r'\{[\s\S]*\}', stripped)
    if not match:
        return None

    try:
        return json.loads(match.group(0))
    except Exception:
        return None


def _normalize_text_list(lines: List[str], limit: int = 3) -> List[str]:
    """清洗模型返回的项目列表。"""
    results = []
    for line in lines:
        cleaned = re.sub(r'^[\-\*\d\.\)\s、]+', '', (line or '').strip())
        cleaned = cleaned.strip('：:;； ')
        if cleaned and cleaned not in results:
            results.append(cleaned)
        if len(results) >= limit:
            break
    return results


def _extract_plaintext_ai_report(content: str, report_context: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """尽量将普通文本回复解析为结构化 AI 报告。"""
    if not content:
        return None

    text = content.strip()
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return None

    section_aliases = {
        'executive_summary': ['执行摘要', '摘要', '总结', '概述'],
        'risk_judgment': ['风险判断', '风险研判', '风险评估', '研判'],
        'action_recommendations': ['处置建议', '行动建议', '建议措施', '建议'],
        'pr_talking_points': ['公关口径', '回应口径', '话术口径', '对外口径']
    }
    sections = {
        'executive_summary': [],
        'risk_judgment': [],
        'action_recommendations': [],
        'pr_talking_points': []
    }

    current_key = None
    for line in lines:
        matched_key = None
        for key, aliases in section_aliases.items():
            if any(line.startswith(alias) for alias in aliases):
                matched_key = key
                break

        if matched_key:
            current_key = matched_key
            trailing = re.split(r'[：:]', line, maxsplit=1)
            if len(trailing) == 2 and trailing[1].strip():
                sections[current_key].append(trailing[1].strip())
            continue

        if current_key:
            sections[current_key].append(line)

    action_items = _normalize_text_list(sections['action_recommendations'])
    pr_items = _normalize_text_list(sections['pr_talking_points'])

    if not action_items:
        numbered_lines = [line for line in lines if re.match(r'^[\d一二三四五六七八九十]+[\.、\)]', line)]
        action_items = _normalize_text_list(numbered_lines)

    paragraph_lines = [line for line in lines if not re.match(r'^[\-\*\d一二三四五六七八九十]+[\.、\)]', line)]
    executive_summary = ' '.join(sections['executive_summary']).strip()
    risk_judgment = ' '.join(sections['risk_judgment']).strip()

    if not executive_summary and paragraph_lines:
        executive_summary = paragraph_lines[0]
    if not risk_judgment and len(paragraph_lines) > 1:
        risk_judgment = paragraph_lines[1]

    if not pr_items and action_items:
        keyword = report_context.get('keyword', '相关问题')
        pr_items = [
            f'我们已关注到与“{keyword}”相关的反馈，并正在核查处理。',
            '对已确认问题将尽快整改，并持续同步进展。',
            '欢迎通过官方渠道反馈，我们会持续优化服务体验。'
        ]

    if not executive_summary and not risk_judgment and not action_items:
        return None

    return {
        'enhanced': True,
        'source': 'ai_text',
        'executive_summary': executive_summary,
        'risk_judgment': risk_judgment,
        'action_recommendations': action_items[:3],
        'pr_talking_points': pr_items[:3]
    }


def _try_generate_ai_report(report_context: Dict[str, Any]) -> Dict[str, Any]:
    """尝试使用兼容 OpenAI 的接口生成结构化 AI 报告，失败则回退规则引擎。"""
    api_key = os.getenv('OPENAI_API_KEY') or os.getenv('AI_API_KEY')
    if not api_key:
        return _build_rule_based_ai_report(report_context)

    model = os.getenv('AI_MODEL', 'gpt-4o-mini')
    base_url = os.getenv('AI_API_BASE', 'https://api.openai.com/v1').rstrip('/')
    endpoint = f"{base_url}/chat/completions"

    prompt = (
        "你是中文舆情和客诉分析顾问。请严格输出 JSON 对象，不要输出 markdown。"
        "JSON 必须包含以下字段：executive_summary(string), risk_judgment(string), "
        "action_recommendations(array of string, 3 items), pr_talking_points(array of string, 3 items)。"
        "内容要聚焦事件研判、业务处置、品牌公关口径，语言专业、简洁、可执行。"
    )

    payload = {
        'model': model,
        'messages': [
            {'role': 'system', 'content': prompt},
            {'role': 'user', 'content': json.dumps(report_context, ensure_ascii=False)}
        ],
        'temperature': 0.3,
        'max_tokens': 320,
        'response_format': {'type': 'json_object'}
    }

    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json'
    }

    try:
        response = requests.post(endpoint, headers=headers, json=payload, timeout=25)
        response.raise_for_status()
        data = response.json()
        content = (
            data.get('choices', [{}])[0]
            .get('message', {})
            .get('content', '')
            .strip()
        )
        parsed = _extract_json_object(content)
        if not parsed:
            plain_report = _extract_plaintext_ai_report(content, report_context)
            if plain_report:
                return plain_report
            logger.warning('AI report returned non-JSON content, fallback to rule-based report')
            return _build_rule_based_ai_report(report_context)

        return {
            'enhanced': True,
            'source': 'ai',
            'executive_summary': parsed.get('executive_summary', ''),
            'risk_judgment': parsed.get('risk_judgment', ''),
            'action_recommendations': list(parsed.get('action_recommendations', []))[:3],
            'pr_talking_points': list(parsed.get('pr_talking_points', []))[:3]
        }
    except Exception as exc:
        logger.warning(f"AI advice generation skipped: {exc}")
        return _build_rule_based_ai_report(report_context)


def _try_generate_ai_advice(report_context: Dict[str, Any]) -> Optional[str]:
    """兼容旧逻辑：返回合并后的 AI 建议字符串。"""
    ai_report = _try_generate_ai_report(report_context)
    action_recommendations = ai_report.get('action_recommendations', [])
    if action_recommendations:
        return ' '.join(action_recommendations)
    return ai_report.get('executive_summary') or None


def _emit_monitor_alert(profile: MonitorProfile, keyword: str, pipeline_result: Dict[str, Any], reason: str, level: AlertLevel) -> None:
    """为监控对象生成自定义预警。"""
    summary = pipeline_result.get('report', {}).get('summary', '')
    alert = Alert(
        alert_id=f"monitor_{profile.monitor_id}_{uuid.uuid4().hex[:10]}",
        rule_id=f"monitor_profile_{profile.monitor_id}",
        level=level,
        title=f"[{profile.name}] 监控阈值触发",
        description=reason,
        source=f"monitor:{profile.name}",
        content=summary,
        matched_keywords=[keyword],
        metadata={
            'monitor_id': profile.monitor_id,
            'keyword': keyword,
            'pipeline_id': pipeline_result.get('pipeline_id')
        },
        status=AlertStatus.ACTIVE
    )
    alert_manager.alerts[alert.alert_id] = alert
    alert_manager.alert_history.append(alert)
    alert_manager._alert_queue.put(alert)


def _evaluate_monitor_thresholds(profile: MonitorProfile, keyword: str, pipeline_result: Dict[str, Any]) -> List[str]:
    """检查监控对象阈值并触发预警。"""
    thresholds = profile.thresholds or {}
    analysis = pipeline_result.get('analysis', {})
    sentiment_stats = analysis.get('sentiment', {}).get('data', {}).get('statistics', {})
    distribution = sentiment_stats.get('distribution', {})
    risk = analysis.get('risk', {}).get('data', {})
    total_items = pipeline_result.get('total_items', 0)

    triggered = []
    negative_ratio = float(distribution.get('negative_ratio', 0) or 0)
    risk_score = float(risk.get('risk_score', 0) or 0)
    if thresholds.get('negative_ratio') is not None and negative_ratio >= float(thresholds['negative_ratio']):
        reason = f"负面占比达到 {negative_ratio:.2%}，超过阈值 {float(thresholds['negative_ratio']):.2%}"
        _emit_monitor_alert(profile, keyword, pipeline_result, reason, AlertLevel.WARNING)
        triggered.append(reason)
    if thresholds.get('risk_score') is not None and risk_score >= float(thresholds['risk_score']):
        reason = f"风险分达到 {risk_score:.1f}，超过阈值 {float(thresholds['risk_score']):.1f}"
        _emit_monitor_alert(profile, keyword, pipeline_result, reason, AlertLevel.DANGER)
        triggered.append(reason)
    if thresholds.get('min_items') is not None and total_items >= int(thresholds['min_items']):
        reason = f"采集量达到 {total_items}，超过监控阈值 {int(thresholds['min_items'])}"
        _emit_monitor_alert(profile, keyword, pipeline_result, reason, AlertLevel.INFO)
        triggered.append(reason)
    return triggered


def _run_dashboard_pipeline_internal(keyword: str, platforms: List[str], max_items: int) -> Dict[str, Any]:
    """复用的一体化分析流水线，供 API 与监控任务调用。"""
    all_items: List[Dict[str, Any]] = []
    platform_stats: Dict[str, Dict[str, Any]] = {}
    errors = []
    collection_started_at = time.perf_counter()

    def collect_platform(platform: str) -> Dict[str, Any]:
        started_at = time.perf_counter()
        collector = CollectorFactory.create(platform)
        items = collector.collect(keyword, limit=max_items)
        normalized_items = []
        collected_at = datetime.now().isoformat()
        for item in items:
            normalized_items.append({
                'id': str(uuid.uuid4()),
                'platform': item.platform,
                'content': item.content,
                'author': item.author,
                'author_id': item.author_id,
                'url': item.url,
                'publish_time': item.publish_time.isoformat() if item.publish_time else None,
                'likes': item.likes,
                'comments': item.comments,
                'shares': item.shares,
                'metadata': item.metadata,
                'collected_at': collected_at
            })
        duration_seconds = round(time.perf_counter() - started_at, 3)
        return {
            'platform': platform,
            'items': normalized_items,
            'duration_seconds': duration_seconds,
        }

    max_workers = max(1, min(len(platforms), 8))
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_map = {executor.submit(collect_platform, platform): platform for platform in platforms}
        for future in as_completed(future_map):
            platform = future_map[future]
            try:
                result = future.result()
                normalized_items = result['items']
                all_items.extend(normalized_items)
                platform_stats[platform] = {
                    'success': True,
                    'items': len(normalized_items),
                    'duration_seconds': result['duration_seconds']
                }
            except Exception as exc:
                error_message = str(exc)
                errors.append({'platform': platform, 'error': error_message})
                platform_stats[platform] = {
                    'success': False,
                    'items': 0,
                    'error': error_message
                }

    # 检查是否任何平台采集成功（无论是否返回数据）
    any_platform_succeeded = any(
        isinstance(ps, dict) and ps.get('success', False) 
        for ps in platform_stats.values()
    )
    
    # 如果没有任何平台成功，返回失败
    if not any_platform_succeeded:
        return {
            'success': False,
            'message': '所有渠道采集失败，未获取有效数据',
            'keyword': keyword,
            'platforms': platforms,
            'platform_stats': platform_stats,
            'errors': errors
        }

    collection_id = str(uuid.uuid4())
    collected_data_store[collection_id] = {
        'platform': 'multi',
        'keyword': keyword,
        'platforms': platforms,
        'collected_at': datetime.now().isoformat(),
        'items': all_items
    }

    # 如果没有采集到数据，生成诊断报告而不是跳过
    if not all_items:
        analysis = {
            'sentiment': {'positive': 0, 'neutral': 0, 'negative': 0},
            'trend': {},
            'risk': {'level': '低', 'score': 0, 'factors': []},
            'relationship': {},
            'summary': {
                'total_items': 0,
                'analysis_time': datetime.now().isoformat(),
                'note': f'关键词"{keyword}"在{len([p for p in platform_stats.values() if isinstance(p, dict) and p.get("success")])}个平台采集成功但无相关数据'
            }
        }
        wordcloud = []
        ai_report = f'关键词"{keyword}"暂无对应舆情数据。建议：1. 检查关键词拼写；2. 尝试相关词汇；3. 联系数据提供商补充覆盖。'
    else:
        analysis = analyzer.comprehensive_analysis(all_items)
        wordcloud = _generate_wordcloud_from_items(all_items)
        ai_report = _try_generate_ai_report({
            'keyword': keyword,
            'platforms': platforms,
            'platform_stats': platform_stats,
            'analysis_summary': analysis.get('summary', {}),
            'risk': analysis.get('risk', {}),
            'wordcloud': wordcloud
        })

    analysis_id = str(uuid.uuid4())
    analysis_results_store[analysis_id] = {
        'created_at': datetime.now().isoformat(),
        'source_count': len(all_items),
        'results': analysis
    }

    report_data = _build_report_text(
        keyword=keyword,
        platforms=platforms,
        total_items=len(all_items),
        analysis=analysis,
        wordcloud=wordcloud,
        ai_report=ai_report
    )

    pipeline_id = str(uuid.uuid4())
    dashboard_pipeline_store[pipeline_id] = {
        'created_at': datetime.now().isoformat(),
        'keyword': keyword,
        'platforms': platforms,
        'collection_id': collection_id,
        'analysis_id': analysis_id,
        'report': report_data,
        'wordcloud': wordcloud,
        'platform_stats': platform_stats,
        'collection_duration_seconds': round(time.perf_counter() - collection_started_at, 3)
    }

    return {
        'success': True,
        'pipeline_id': pipeline_id,
        'keyword': keyword,
        'platforms': platforms,
        'collection_id': collection_id,
        'analysis_id': analysis_id,
        'platform_stats': platform_stats,
        'collection_duration_seconds': round(time.perf_counter() - collection_started_at, 3),
        'total_items': len(all_items),
        'analysis': analysis,
        'wordcloud': wordcloud,
        'report': report_data,
        'errors': errors
    }


def _schedule_monitor_profile(profile: MonitorProfile) -> None:
    """将监控对象注册到 APScheduler。"""
    if not getattr(scheduler, '_scheduler', None) or not profile.enabled:
        return
    job_id = f"monitor_profile_{profile.monitor_id}"
    try:
        scheduler._scheduler.remove_job(job_id)
    except Exception:
        pass
    scheduler._scheduler.add_job(
        func=_run_monitor_profile_once,
        trigger='interval',
        seconds=max(60, int(profile.interval_seconds)),
        id=job_id,
        replace_existing=True,
        args=[profile.monitor_id]
    )
    profile.scheduler_job_id = job_id


def _unschedule_monitor_profile(monitor_id: str) -> None:
    job_id = f"monitor_profile_{monitor_id}"
    if not getattr(scheduler, '_scheduler', None):
        return
    try:
        scheduler._scheduler.remove_job(job_id)
    except Exception:
        pass


def _run_monitor_profile_once(monitor_id: str) -> Dict[str, Any]:
    """执行单个监控对象。"""
    profile = monitor_profiles_store.get(monitor_id)
    if not profile or not profile.enabled:
        return {'success': False, 'message': 'monitor disabled or missing'}

    pipeline_ids = []
    triggered_reasons = []
    last_result = None
    for keyword in profile.keywords:
        result = _run_dashboard_pipeline_internal(keyword, profile.platforms, profile.max_items)
        last_result = result
        if result.get('success'):
            pipeline_ids.append(result.get('pipeline_id'))
            triggered_reasons.extend(_evaluate_monitor_thresholds(profile, keyword, result))

    profile.last_run_at = datetime.now().isoformat()
    profile.updated_at = profile.last_run_at
    profile.last_pipeline_ids = pipeline_ids[-10:]
    profile.last_status = 'success' if pipeline_ids else 'failed'

    return {
        'success': bool(pipeline_ids),
        'monitor_id': monitor_id,
        'pipeline_ids': pipeline_ids,
        'triggered_reasons': triggered_reasons,
        'last_result': last_result
    }


def _export_report_file(pipeline_id: str, export_format: str) -> Path:
    """导出报告文件。"""
    pipeline = dashboard_pipeline_store.get(pipeline_id)
    if not pipeline:
        raise APIError('Pipeline not found', 404)

    report = pipeline.get('report', {})
    analysis_record = analysis_results_store.get(pipeline.get('analysis_id', ''), {})
    analysis = analysis_record.get('results', {})
    sentiment_stats = analysis.get('sentiment', {}).get('data', {}).get('statistics', {})
    risk_data = analysis.get('risk', {}).get('data', {})
    trend_data = analysis.get('trend', {}).get('data', {})
    wordcloud = pipeline.get('wordcloud', [])
    hot_words = '、'.join(item.get('name') for item in wordcloud[:10] if item.get('name')) or '暂无'
    platform_stats = pipeline.get('platform_stats', {})
    group_overview = _build_group_overview_payload()
    competitor_overview = _build_competitor_overview_payload()
    export_dir = _ensure_export_dir()
    safe_keyword = re.sub(r'[^\w\u4e00-\u9fa5-]+', '_', pipeline.get('keyword', 'report'))
    filename = f"{safe_keyword}_{pipeline_id[:8]}.{export_format}"
    output_path = export_dir / filename

    if export_format == 'docx':
        from docx import Document
        document = Document()
        document.add_heading(report.get('title', '舆情分析报告'), 0)
        document.add_paragraph('OSINT CN 智能舆情分析报告')
        document.add_paragraph(f"生成时间：{report.get('generated_at', '-')}")
        document.add_paragraph(f"监测关键词：{pipeline.get('keyword', '-')}")
        document.add_paragraph(f"监测平台：{', '.join(pipeline.get('platforms', []))}")
        document.add_page_break()
        document.add_heading('一、执行摘要', level=1)
        document.add_paragraph(report.get('summary', '-'))
        document.add_heading('二、核心指标', level=1)
        document.add_paragraph(f"采集总量：{analysis.get('summary', {}).get('total_items', 0)}")
        document.add_paragraph(f"风险等级：{risk_data.get('risk_level', '-')}")
        document.add_paragraph(f"风险分：{risk_data.get('risk_score', 0)}")
        document.add_paragraph(
            f"情绪分布：正面 {sentiment_stats.get('positive_count', 0)} / 中立 {sentiment_stats.get('neutral_count', 0)} / 负面 {sentiment_stats.get('negative_count', 0)}"
        )
        document.add_paragraph(f"传播峰值时段：{trend_data.get('peak_time', '-')}")
        document.add_heading('摘要', level=1)
        document.add_paragraph(report.get('summary', '-'))
        document.add_heading('关键发现', level=1)
        for item in report.get('findings', []):
            document.add_paragraph(item, style='List Bullet')
        document.add_heading('处置建议', level=1)
        for item in report.get('recommendations', []):
            document.add_paragraph(item, style='List Bullet')
        ai_insight = report.get('ai_insight', {})
        document.add_heading('AI 研判', level=1)
        document.add_paragraph(ai_insight.get('executive_summary', '-'))
        document.add_paragraph(ai_insight.get('risk_judgment', '-'))
        document.add_heading('AI 公关口径', level=2)
        for item in ai_insight.get('pr_talking_points', []):
            document.add_paragraph(item, style='List Bullet')
        document.add_heading('附录', level=1)
        document.add_paragraph(f"高频热词：{hot_words}")
        if platform_stats:
            document.add_paragraph('平台采集明细：')
            for platform, stats in platform_stats.items():
                document.add_paragraph(f"{platform}: {stats.get('items', 0)} 条", style='List Bullet')
        if group_overview.get('groups'):
            document.add_heading('监控分组对比', level=2)
            for item in group_overview.get('groups', [])[:5]:
                document.add_paragraph(
                    f"{item.get('name', '-')}: 热度 {item.get('heat', 0)}，风险分 {item.get('risk_score', 0)}，负面占比 {int((item.get('negative_ratio', 0) or 0) * 100)}%",
                    style='List Bullet'
                )
        if competitor_overview.get('comparisons'):
            document.add_heading('竞品对比摘要', level=2)
            base_group = competitor_overview.get('base_group', {})
            document.add_paragraph(f"对比基准：{base_group.get('name', '暂无')}")
            for item in competitor_overview.get('comparisons', [])[:5]:
                document.add_paragraph(
                    f"对比 {item.get('rival_name', '-')}: 结论 {item.get('status', '-') }，热度差 {item.get('heat_gap', 0)}，风险分差 {item.get('risk_gap', 0)}，共同关注 {('、'.join(item.get('keyword_overlap', [])) or '暂无')}",
                    style='List Bullet'
                )
        document.save(str(output_path))
        return output_path

    if export_format == 'pdf':
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        font_candidates = [
            '/System/Library/Fonts/PingFang.ttc',
            '/System/Library/Fonts/STHeiti Light.ttc',
            '/Library/Fonts/Arial Unicode.ttf'
        ]
        registered_font = 'Helvetica'
        for candidate in font_candidates:
            if Path(candidate).exists():
                try:
                    pdfmetrics.registerFont(TTFont('CustomCJK', candidate))
                    registered_font = 'CustomCJK'
                    break
                except Exception:
                    continue

        styles = getSampleStyleSheet()
        for style in styles.byName.values():
            style.fontName = registered_font
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        story = [
            Paragraph(report.get('title', '舆情分析报告'), styles['Title']),
            Spacer(1, 12),
            Paragraph('OSINT CN 智能舆情分析报告', styles['Heading2']),
            Spacer(1, 12),
            Paragraph(f"生成时间：{report.get('generated_at', '-')}", styles['Normal']),
            Paragraph(f"关键词：{pipeline.get('keyword', '-')}", styles['Normal']),
            Paragraph(f"平台：{', '.join(pipeline.get('platforms', []))}", styles['Normal']),
            Spacer(1, 12),
            Paragraph('执行摘要', styles['Heading2']),
            Paragraph(report.get('summary', '-'), styles['Normal']),
            Spacer(1, 12),
            Paragraph('核心指标', styles['Heading2']),
            Paragraph(f"采集总量：{analysis.get('summary', {}).get('total_items', 0)}", styles['Normal']),
            Paragraph(f"风险等级：{risk_data.get('risk_level', '-')} / 风险分：{risk_data.get('risk_score', 0)}", styles['Normal']),
            Paragraph(
                f"情绪分布：正面 {sentiment_stats.get('positive_count', 0)} / 中立 {sentiment_stats.get('neutral_count', 0)} / 负面 {sentiment_stats.get('negative_count', 0)}",
                styles['Normal']
            ),
            Paragraph(f"传播峰值时段：{trend_data.get('peak_time', '-')}", styles['Normal']),
            Spacer(1, 12),
            Paragraph('关键发现', styles['Heading2'])
        ]
        for item in report.get('findings', []):
            story.append(Paragraph(f"• {item}", styles['Normal']))
        story.append(Spacer(1, 12))
        story.append(Paragraph('处置建议', styles['Heading2']))
        for item in report.get('recommendations', []):
            story.append(Paragraph(f"• {item}", styles['Normal']))
        ai_insight = report.get('ai_insight', {})
        story.append(Spacer(1, 12))
        story.append(Paragraph('AI 研判', styles['Heading2']))
        story.append(Paragraph(ai_insight.get('executive_summary', '-'), styles['Normal']))
        story.append(Paragraph(ai_insight.get('risk_judgment', '-'), styles['Normal']))
        for item in ai_insight.get('pr_talking_points', []):
            story.append(Paragraph(f"• {item}", styles['Normal']))
        story.append(Spacer(1, 12))
        story.append(Paragraph('附录', styles['Heading2']))
        story.append(Paragraph(f"高频热词：{hot_words}", styles['Normal']))
        for platform, stats in platform_stats.items():
            story.append(Paragraph(f"• {platform}: {stats.get('items', 0)} 条", styles['Normal']))
        if group_overview.get('groups'):
            story.append(Spacer(1, 12))
            story.append(Paragraph('监控分组对比', styles['Heading2']))
            for item in group_overview.get('groups', [])[:5]:
                story.append(Paragraph(
                    f"• {item.get('name', '-')}: 热度 {item.get('heat', 0)}，风险分 {item.get('risk_score', 0)}，负面占比 {int((item.get('negative_ratio', 0) or 0) * 100)}%",
                    styles['Normal']
                ))
        if competitor_overview.get('comparisons'):
            story.append(Spacer(1, 12))
            story.append(Paragraph('竞品对比摘要', styles['Heading2']))
            base_group = competitor_overview.get('base_group', {})
            story.append(Paragraph(f"对比基准：{base_group.get('name', '暂无')}", styles['Normal']))
            for item in competitor_overview.get('comparisons', [])[:5]:
                story.append(Paragraph(
                    f"• 对比 {item.get('rival_name', '-')}: 结论 {item.get('status', '-')}，热度差 {item.get('heat_gap', 0)}，风险分差 {item.get('risk_gap', 0)}，共同关注 {('、'.join(item.get('keyword_overlap', [])) or '暂无')}",
                    styles['Normal']
                ))
        doc.build(story)
        return output_path

    if export_format == 'html':
        keyword = pipeline.get('keyword', '舆情报告')
        platforms = pipeline.get('platforms', [])
        total_items = analysis.get('summary', {}).get('total_items', 0)
        analysis_for_html = analysis
        wordcloud = pipeline.get('wordcloud', [])
        ai_report = report.get('ai_insight', {})
        
        html_content = _build_report_html(keyword, platforms, total_items, analysis_for_html, wordcloud, ai_report)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        return output_path

    raise APIError('Unsupported export format', 400)


def _build_group_overview_payload() -> Dict[str, Any]:
    """构建监控分组总览数据。"""
    group_profiles: Dict[str, Dict[str, Any]] = {}

    def ensure_group(group_id: str, name: str, color: str = '#6fcbff') -> Dict[str, Any]:
        if group_id not in group_profiles:
            group_profiles[group_id] = {
                'group_id': group_id,
                'name': name,
                'color': color,
                'monitors': [],
                'pipeline_ids': [],
                'heat': 0,
                'risk_scores': [],
                'negative_ratios': [],
                'word_counter': Counter(),
                'latest_summary': '',
                'latest_time': ''
            }
        return group_profiles[group_id]

    group_lookup = {group_id: group for group_id, group in monitor_groups_store.items()}
    for profile in monitor_profiles_store.values():
        group_id = profile.group_id or 'ungrouped'
        group_name = group_lookup[group_id].name if group_id in group_lookup else '未分组'
        group_color = group_lookup[group_id].color if group_id in group_lookup else '#6fcbff'
        bucket = ensure_group(group_id, group_name, group_color)
        bucket['monitors'].append(profile.name)
        bucket['pipeline_ids'].extend(profile.last_pipeline_ids)

    group_items = []
    trend = {}
    alerts = []
    for group_id, bucket in group_profiles.items():
        unique_pipeline_ids = []
        seen = set()
        for pipeline_id in bucket['pipeline_ids']:
            if pipeline_id and pipeline_id not in seen:
                unique_pipeline_ids.append(pipeline_id)
                seen.add(pipeline_id)

        for pipeline_id in unique_pipeline_ids:
            pipeline = dashboard_pipeline_store.get(pipeline_id)
            if not pipeline:
                continue
            analysis = analysis_results_store.get(pipeline.get('analysis_id', ''), {}).get('results', {})
            sentiment_stats = analysis.get('sentiment', {}).get('data', {}).get('statistics', {})
            distribution = sentiment_stats.get('distribution', {})
            risk_data = analysis.get('risk', {}).get('data', {})
            bucket['heat'] += int(analysis.get('summary', {}).get('total_items', 0) or 0)
            bucket['risk_scores'].append(float(risk_data.get('risk_score', 0) or 0))
            bucket['negative_ratios'].append(float(distribution.get('negative_ratio', 0) or 0))
            for item in pipeline.get('wordcloud', [])[:15]:
                name = item.get('name')
                value = int(item.get('value', 0) or 0)
                if name:
                    bucket['word_counter'][name] += value
            bucket['latest_summary'] = pipeline.get('report', {}).get('summary', bucket['latest_summary'])
            bucket['latest_time'] = max(bucket['latest_time'], pipeline.get('created_at', '') or '')

        risk_avg = round(sum(bucket['risk_scores']) / len(bucket['risk_scores']), 2) if bucket['risk_scores'] else 0
        negative_avg = round(sum(bucket['negative_ratios']) / len(bucket['negative_ratios']), 3) if bucket['negative_ratios'] else 0
        top_keywords = [name for name, _ in bucket['word_counter'].most_common(6)]
        item = {
            'group_id': group_id,
            'name': bucket['name'],
            'color': bucket['color'],
            'monitor_count': len(bucket['monitors']),
            'pipeline_count': len(unique_pipeline_ids),
            'heat': bucket['heat'],
            'risk_score': risk_avg,
            'negative_ratio': negative_avg,
            'top_keywords': top_keywords,
            'latest_summary': bucket['latest_summary'],
            'latest_time': bucket['latest_time']
        }
        group_items.append(item)
        trend[bucket['name']] = bucket['heat']
        if risk_avg >= 50 or negative_avg >= 0.3:
            alerts.append({
                'message': f"{bucket['name']} 风险分 {risk_avg}，负面占比 {negative_avg:.1%}，建议优先关注。"
            })

    group_items.sort(key=lambda item: (item['risk_score'], item['heat']), reverse=True)
    top_group = group_items[0] if group_items else None
    news = [
        f"{item['name']}：热度 {item['heat']}，风险分 {item['risk_score']}，关注词 {('、'.join(item['top_keywords']) or '暂无')}"
        for item in group_items[:8]
    ]

    return {
        'success': True,
        'groups': group_items,
        'overview': {
            'group_count': len(group_items),
            'monitor_count': len(monitor_profiles_store),
            'top_group': top_group['name'] if top_group else '暂无',
            'top_risk_score': top_group['risk_score'] if top_group else 0
        },
        'trend': trend,
        'news': news,
        'alerts': alerts,
        'hot_words': top_group['top_keywords'] if top_group else [],
        'detail': top_group or {}
    }


def _build_competitor_overview_payload(base_group_id: Optional[str] = None) -> Dict[str, Any]:
    """构建监控分组竞品对比数据。"""
    overview_payload = _build_group_overview_payload()
    groups = overview_payload.get('groups', [])

    if not groups:
        return {
            'success': True,
            'groups': [],
            'base_group': {},
            'comparisons': [],
            'overview': {
                'base_group': '暂无',
                'rival_count': 0,
                'strongest_rival': '暂无'
            },
            'trend': {},
            'news': ['暂无竞品对比数据，请先创建并执行至少两个监控分组。'],
            'alerts': [],
            'hot_words': [],
            'detail': {}
        }

    base_group = next((item for item in groups if item.get('group_id') == base_group_id), None)
    if not base_group:
        base_group = groups[0]

    comparisons = []
    for rival in groups:
        if rival.get('group_id') == base_group.get('group_id'):
            continue

        heat_gap = int((base_group.get('heat', 0) or 0) - (rival.get('heat', 0) or 0))
        risk_gap = round((base_group.get('risk_score', 0) or 0) - (rival.get('risk_score', 0) or 0), 2)
        negative_gap = round((base_group.get('negative_ratio', 0) or 0) - (rival.get('negative_ratio', 0) or 0), 3)
        keyword_overlap = sorted(set(base_group.get('top_keywords', [])) & set(rival.get('top_keywords', [])))

        if heat_gap >= 0 and risk_gap <= 0:
            status = '领先'
            summary = f"{base_group.get('name', '基准组')} 当前热度更高且风险更低。"
        elif heat_gap < 0 and risk_gap > 0:
            status = '承压'
            summary = f"{rival.get('name', '竞品组')} 热度更高且风险更低，建议优先补位。"
        else:
            status = '胶着'
            summary = f"{base_group.get('name', '基准组')} 与 {rival.get('name', '竞品组')} 处于拉锯状态。"

        comparisons.append({
            'rival_group_id': rival.get('group_id'),
            'rival_name': rival.get('name'),
            'status': status,
            'summary': summary,
            'heat_gap': heat_gap,
            'risk_gap': risk_gap,
            'negative_gap': negative_gap,
            'base_heat': base_group.get('heat', 0),
            'rival_heat': rival.get('heat', 0),
            'base_risk_score': base_group.get('risk_score', 0),
            'rival_risk_score': rival.get('risk_score', 0),
            'base_negative_ratio': base_group.get('negative_ratio', 0),
            'rival_negative_ratio': rival.get('negative_ratio', 0),
            'keyword_overlap': keyword_overlap,
            'rival_keywords': rival.get('top_keywords', [])
        })

    comparisons.sort(key=lambda item: (item['rival_heat'], -item['rival_risk_score']), reverse=True)
    strongest_rival = comparisons[0] if comparisons else None
    alerts = []
    if strongest_rival and strongest_rival.get('status') == '承压':
        alerts.append({
            'message': f"{base_group.get('name', '基准组')} 相比 {strongest_rival.get('rival_name', '竞品组')} 处于承压状态，建议优先优化共同关注词相关议题。"
        })

    news = [
        f"{item.get('rival_name', '竞品组')}：{item.get('summary', '-') } 共同关注 {('、'.join(item.get('keyword_overlap', [])) or '暂无')}"
        for item in comparisons[:8]
    ] or ['当前仅有一个监控分组，暂无法形成竞品对比。']

    trend = {base_group.get('name', '基准组'): base_group.get('heat', 0)}
    for item in comparisons[:6]:
        trend[item.get('rival_name', '竞品组')] = item.get('rival_heat', 0)

    return {
        'success': True,
        'groups': groups,
        'base_group': base_group,
        'comparisons': comparisons,
        'overview': {
            'base_group': base_group.get('name', '暂无'),
            'rival_count': len(comparisons),
            'strongest_rival': strongest_rival.get('rival_name', '暂无') if strongest_rival else '暂无'
        },
        'trend': trend,
        'news': news,
        'alerts': alerts,
        'hot_words': base_group.get('top_keywords', []),
        'detail': strongest_rival or {}
    }


# ============ 错误处理 ============

class APIError(Exception):
    """API 错误"""
    def __init__(self, message: str, status_code: int = 400):
        self.message = message
        self.status_code = status_code


@app.errorhandler(APIError)
def handle_api_error(error):
    return jsonify({'error': error.message, 'status': 'error'}), error.status_code


@app.errorhandler(404)
def not_found(error):
    return jsonify({
        'error': 'Not Found',
        'message': '请求的资源不存在',
        'available_endpoints': [
            'GET /',
            'GET /health',
            'GET /api/platforms',
            'POST /api/collect',
            'POST /api/analyze',
            'POST /api/process-text',
            'POST /api/sentiment',
            'POST /api/keywords',
            'GET /api/report'
        ]
    }), 404


@app.errorhandler(500)
def internal_error(error):
    logger.error(f"Internal error: {error}")
    return jsonify({'error': 'Internal Server Error', 'status': 'error'}), 500


def require_json(f):
    """装饰器：要求 JSON 请求体"""
    @wraps(f)
    def decorated(*args, **kwargs):
        if not request.is_json:
            raise APIError('Content-Type must be application/json', 400)
        return f(*args, **kwargs)
    return decorated


def validate_request(model_class):
    """装饰器：使用 Pydantic 模型验证请求"""
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            try:
                data = request.get_json() or {}
                validated = model_class(**data)
                g.validated_data = validated
                return f(*args, **kwargs)
            except ValidationError as e:
                errors = []
                for err in e.errors():
                    errors.append({
                        'field': '.'.join(str(x) for x in err['loc']),
                        'message': err['msg'],
                        'type': err['type']
                    })
                return jsonify({
                    'success': False,
                    'error': 'validation_error',
                    'message': '请求参数验证失败',
                    'details': errors
                }), 400
        return decorated
    return decorator


PUBLIC_PATH_PREFIXES = ('/', '/dashboard', '/health')
PUBLIC_API_PATHS = {'/api/platforms'}


@app.before_request
def enforce_api_key_if_needed():
    if not ENFORCE_API_KEY:
        return None

    path = request.path or '/'
    if path in PUBLIC_PATH_PREFIXES or path in PUBLIC_API_PATHS:
        return None

    if path.startswith('/api/'):
        api_key = request.headers.get('X-API-Key') or request.args.get('api_key')
        if not api_key:
            return jsonify({
                'success': False,
                'error': 'unauthorized',
                'message': '生产模式已开启 API Key 强制校验，请提供 X-API-Key'
            }), 401

        key_info = api_key_manager.validate(api_key)
        if not key_info:
            return jsonify({
                'success': False,
                'error': 'forbidden',
                'message': 'API Key 无效或已禁用'
            }), 403

        g.api_key = api_key
        g.api_key_info = key_info

    return None


# ============ 首页和健康检查 ============

# 导入仪表板模板
from osint_cn.dashboard import DASHBOARD_HTML

INDEX_HTML = '''
<!DOCTYPE html>
<html>
<head>
    <title>OSINT CN - 开源情报系统</title>
    <meta charset="utf-8">
    <meta http-equiv="refresh" content="0; url=/dashboard">
    <style>
        body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; 
               max-width: 900px; margin: 50px auto; padding: 20px; background: #f5f5f5; text-align: center; }
        a { color: #007bff; text-decoration: none; font-size: 1.2rem; }
    </style>
</head>
<body>
    <h1>🔍 OSINT CN</h1>
    <p>正在跳转到仪表板...</p>
    <p><a href="/dashboard">点击这里手动跳转</a></p>
</body>
</html>
'''

START_TIME = datetime.now()


@app.route('/')
def index():
    """首页"""
    return render_template_string(INDEX_HTML, start_time=START_TIME.strftime('%Y-%m-%d %H:%M:%S'))


@app.route('/dashboard')
def dashboard():
    """仪表板页面"""
    return render_template_string(DASHBOARD_HTML)


@app.route('/health', methods=['GET'])
@rate_limit(limit=100)
def health():
    """健康检查"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'version': '1.0.0',
        'uptime_seconds': (datetime.now() - START_TIME).total_seconds(),
        'scheduler': scheduler.get_scheduler_info()
    })


# ============ 平台信息 ============

@app.route('/api/platforms', methods=['GET'])
@rate_limit(limit=100)
def get_platforms():
    """获取支持的平台列表"""
    platforms_info = CollectorFactory.get_platform_info()
    return jsonify({
        'success': True,
        'platforms': platforms_info,
        'count': len(platforms_info)
    })


# ============ 数据采集 ============

@app.route('/api/collect', methods=['POST'])
@require_json
@optional_api_key
@rate_limit(limit=30)
def collect():
    """
    数据采集接口
    
    Request:
        {
            "platform": "weibo",
            "keyword": "搜索关键词",
            "max_items": 100
        }
    """
    data = request.json
    
    # 使用 Pydantic 验证
    try:
        req = CollectRequest(**data)
    except ValidationError as e:
        return jsonify({
            'success': False,
            'error': 'validation_error',
            'details': e.errors()
        }), 400
    
    try:
        collector = CollectorFactory.create(req.platform.value)
        items = collector.collect(req.keyword, limit=req.max_items)
        
        # 生成采集 ID 并存储
        collection_id = str(uuid.uuid4())
        
        # 转换为可序列化格式
        results = []
        for item in items:
            item_data = {
                'id': str(uuid.uuid4()),
                'platform': item.platform,
                'content': item.content,
                'author': item.author,
                'author_id': item.author_id,
                'url': item.url,
                'publish_time': item.publish_time.isoformat() if item.publish_time else None,
                'likes': item.likes,
                'comments': item.comments,
                'shares': item.shares,
                'metadata': item.metadata
            }
            results.append(item_data)
        
        # 存储采集数据
        collected_data_store[collection_id] = {
            'platform': req.platform.value,
            'keyword': req.keyword,
            'collected_at': datetime.now().isoformat(),
            'items': results
        }

        try:
            collection_record = CollectionRecord(
                id=collection_id,
                platform=req.platform.value,
                keyword=req.keyword,
                collected_at=datetime.now(),
                items_count=len(results),
                metadata={'api_source': 'collect'}
            )
            storage_backend.save_collection(collection_record)

            item_records: List[CollectedItemRecord] = []
            for item_data in results:
                publish_time = None
                if item_data.get('publish_time'):
                    try:
                        publish_time = datetime.fromisoformat(item_data['publish_time'])
                    except Exception:
                        publish_time = None

                item_records.append(CollectedItemRecord(
                    id=item_data['id'],
                    collection_id=collection_id,
                    platform=item_data['platform'],
                    content=item_data['content'],
                    author=item_data.get('author', ''),
                    author_id=item_data.get('author_id', ''),
                    url=item_data.get('url', ''),
                    publish_time=publish_time,
                    likes=item_data.get('likes', 0),
                    comments=item_data.get('comments', 0),
                    shares=item_data.get('shares', 0),
                    metadata=item_data.get('metadata', {})
                ))

            if item_records:
                storage_backend.save_items(collection_id, item_records)
        except Exception as e:
            logger.warning(f"Storage persistence failed, fallback to memory only: {e}")
        
        return jsonify({
            'success': True,
            'collection_id': collection_id,
            'platform': req.platform.value,
            'keyword': req.keyword,
            'items_collected': len(results),
            'data': results[:20]  # 只返回前 20 条
        }), 200
        
    except ValueError as e:
        raise APIError(str(e), 400)
    except Exception as e:
        logger.error(f"Collection error: {e}")
        raise APIError(f'Collection failed: {str(e)}', 500)


@app.route('/api/collections', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def list_collections():
    """获取采集记录列表"""
    page = request.args.get('page', 1, type=int)
    page_size = request.args.get('page_size', 20, type=int)

    try:
        records = storage_backend.list_collections(page=page, page_size=page_size)
        collections = [
            {
                'id': r.id,
                'platform': r.platform,
                'keyword': r.keyword,
                'collected_at': r.collected_at.isoformat() if hasattr(r.collected_at, 'isoformat') else r.collected_at,
                'items_count': r.items_count,
                'status': getattr(r, 'status', 'completed')
            }
            for r in records
        ]
        if collections:
            return jsonify({
                'success': True,
                'collections': collections,
                'count': len(collections),
                'page': page,
                'page_size': page_size
            })
    except Exception as e:
        logger.warning(f"Storage list failed, fallback to memory: {e}")

    collections = []
    for cid, data in collected_data_store.items():
        collections.append({
            'id': cid,
            'platform': data['platform'],
            'keyword': data['keyword'],
            'collected_at': data['collected_at'],
            'items_count': len(data['items'])
        })
    
    return jsonify({
        'success': True,
        'collections': collections,
        'count': len(collections)
    })


@app.route('/api/collections/<collection_id>', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_collection(collection_id):
    """获取采集数据详情"""
    page = request.args.get('page', 1, type=int)
    page_size = request.args.get('page_size', 20, type=int)

    try:
        record = storage_backend.get_collection(collection_id)
        if record:
            rows = storage_backend.get_items(collection_id, page=page, page_size=page_size)
            data_rows = []
            for row in rows:
                data_rows.append({
                    'id': row.id,
                    'platform': row.platform,
                    'content': row.content,
                    'author': row.author,
                    'author_id': row.author_id,
                    'url': row.url,
                    'publish_time': row.publish_time.isoformat() if row.publish_time else None,
                    'likes': row.likes,
                    'comments': row.comments,
                    'shares': row.shares,
                    'metadata': row.metadata,
                    'collected_at': row.created_at.isoformat() if row.created_at else None
                })

            return jsonify({
                'success': True,
                'collection_id': record.id,
                'platform': record.platform,
                'keyword': record.keyword,
                'collected_at': record.collected_at.isoformat() if hasattr(record.collected_at, 'isoformat') else record.collected_at,
                'total': record.items_count,
                'page': page,
                'page_size': page_size,
                'data': data_rows
            })
    except Exception as e:
        logger.warning(f"Storage detail failed, fallback to memory: {e}")

    data = collected_data_store.get(collection_id)
    if not data:
        raise APIError('Collection not found', 404)
    
    items = data['items']
    start = (page - 1) * page_size
    end = start + page_size
    
    return jsonify({
        'success': True,
        'collection_id': collection_id,
        'platform': data['platform'],
        'keyword': data['keyword'],
        'collected_at': data['collected_at'],
        'total': len(items),
        'page': page,
        'page_size': page_size,
        'data': items[start:end]
    })


# ============ 文本处理 ============

@app.route('/api/process-text', methods=['POST'])
@require_json
@optional_api_key
@rate_limit(limit=60)
def process_text():
    """
    文本处理接口
    
    Request:
        {
            "text": "要处理的文本",
            "operations": ["segment", "clean", "entities"]
        }
    """
    data = request.json
    text = data.get('text', '')
    operations = data.get('operations', ['segment', 'clean'])
    
    if not text:
        raise APIError('Missing required field: text')
    
    result = {'original': text}
    
    if 'clean' in operations:
        result['cleaned'] = text_processor.clean_text(text)
    
    if 'segment' in operations:
        words = text_processor.segment(text)
        result['segments'] = words
        result['segments_without_stopwords'] = text_processor.remove_stopwords(words)
    
    if 'entities' in operations:
        result['entities'] = text_processor.extract_entities(text)
    
    return jsonify({
        'success': True,
        'result': result
    }), 200


@app.route('/api/sentiment', methods=['POST'])
@require_json
@optional_api_key
@rate_limit(limit=60)
def sentiment():
    """
    情感分析接口
    
    Request:
        {"text": "文本内容"}
        或
        {"texts": ["文本1", "文本2"]}
    """
    data = request.json
    
    if 'text' in data:
        result = text_processor.sentiment_analysis(data['text'])
        return jsonify({
            'success': True,
            'result': {
                'sentiment': result.sentiment,
                'score': result.score,
                'confidence': result.confidence
            }
        }), 200
    
    elif 'texts' in data:
        texts = data['texts']
        if not isinstance(texts, list):
            raise APIError('texts must be a list')
        
        results = []
        for text in texts:
            result = text_processor.sentiment_analysis(text)
            results.append({
                'text': text[:100],
                'sentiment': result.sentiment,
                'score': result.score,
                'confidence': result.confidence
            })
        
        return jsonify({
            'success': True,
            'results': results,
            'count': len(results)
        }), 200
    
    else:
        raise APIError('Missing required field: text or texts')


@app.route('/api/keywords', methods=['POST'])
@require_json
@optional_api_key
@rate_limit(limit=60)
def keywords():
    """
    关键词提取接口
    
    Request:
        {
            "text": "文本内容",
            "top_k": 10,
            "method": "tfidf"  // tfidf 或 textrank
        }
    """
    data = request.json
    text = data.get('text', '')
    top_k = data.get('top_k', 10)
    method = data.get('method', 'tfidf')
    
    if not text:
        raise APIError('Missing required field: text')
    
    keywords = text_processor.extract_keywords(text, top_k=top_k, method=method)
    
    return jsonify({
        'success': True,
        'keywords': [
            {'keyword': kw.keyword, 'weight': kw.weight, 'frequency': kw.frequency}
            for kw in keywords
        ]
    }), 200


# ============ 数据分析 ============

@app.route('/api/analyze', methods=['POST'])
@require_json
@optional_api_key
@rate_limit(limit=30)
def analyze():
    """
    综合分析接口
    
    Request:
        {
            "data": [
                {"content": "文本1", "author": "作者1", ...},
                {"content": "文本2", "author": "作者2", ...}
            ]
        }
        或
        {
            "collection_id": "采集ID"
        }
    """
    req_data = request.json
    
    # 支持直接传入数据或使用采集 ID
    if 'collection_id' in req_data:
        collection_id = req_data['collection_id']
        stored = collected_data_store.get(collection_id)
        if not stored:
            raise APIError('Collection not found', 404)
        data = stored['items']
    else:
        data = req_data.get('data', [])
    
    if not data:
        raise APIError('Missing required field: data or collection_id')
    
    if not isinstance(data, list):
        raise APIError('data must be a list')
    
    # 执行分析
    results = analyzer.comprehensive_analysis(data)
    
    # 存储分析结果
    analysis_id = str(uuid.uuid4())
    analysis_results_store[analysis_id] = {
        'created_at': datetime.now().isoformat(),
        'source_count': len(data),
        'results': results
    }
    
    return jsonify({
        'success': True,
        'analysis_id': analysis_id,
        'analysis': results
    }), 200


@app.route('/api/analyze/sentiment', methods=['POST'])
@require_json
@optional_api_key
@rate_limit(limit=30)
def analyze_sentiment():
    """批量情感分析"""
    texts = request.json.get('texts', [])
    
    if not texts:
        raise APIError('Missing required field: texts')
    
    result = analyzer.sentiment_analysis(texts)
    
    return jsonify({
        'success': True,
        'result': result.to_dict()
    }), 200


@app.route('/api/analyze/trend', methods=['POST'])
@require_json
@optional_api_key
@rate_limit(limit=30)
def analyze_trend():
    """趋势分析"""
    data = request.json.get('data', [])
    time_field = request.json.get('time_field', 'publish_time')
    interval = request.json.get('interval', 'day')
    
    if not data:
        raise APIError('Missing required field: data')
    
    result = analyzer.trend_analysis(data, time_field=time_field, interval=interval)
    
    return jsonify({
        'success': True,
        'result': result.to_dict()
    }), 200


@app.route('/api/analyze/risk', methods=['POST'])
@require_json
@optional_api_key
@rate_limit(limit=30)
def analyze_risk():
    """风险评估"""
    data = request.json.get('data', [])
    context = request.json.get('context', {})
    
    if not data:
        raise APIError('Missing required field: data')
    
    result = analyzer.risk_assessment(data, context=context)
    
    return jsonify({
        'success': True,
        'result': result.to_dict()
    }), 200


# ============ 任务管理 ============

@app.route('/api/tasks', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def list_tasks():
    """获取任务列表"""
    tasks = scheduler.get_all_tasks()
    
    task_list = []
    for task in tasks:
        task_list.append({
            'id': task.id,
            'name': task.name,
            'type': task.task_type.value,
            'status': task.status.value,
            'enabled': task.enabled,
            'schedule': task.schedule,
            'interval_seconds': task.interval_seconds,
            'created_at': task.created_at.isoformat(),
            'last_run': task.last_run.isoformat() if task.last_run else None,
            'next_run': task.next_run.isoformat() if task.next_run else None,
            'run_count': task.run_count
        })
    
    return jsonify({
        'success': True,
        'tasks': task_list,
        'count': len(task_list),
        'scheduler_info': scheduler.get_scheduler_info()
    })


@app.route('/api/tasks', methods=['POST'])
@require_json
@require_api_key
@rate_limit(limit=30)
def create_task():
    """创建任务"""
    data = request.json
    
    name = data.get('name')
    task_type = data.get('task_type', 'collect')
    schedule = data.get('schedule')  # Cron 表达式
    interval_seconds = data.get('interval_seconds')
    config_data = data.get('config', {})
    
    if not name:
        raise APIError('Missing required field: name')
    
    # 解析任务类型
    try:
        task_type_enum = TaskType(task_type)
    except ValueError:
        raise APIError(f'Invalid task_type: {task_type}')
    
    # 创建任务配置
    task_config = TaskConfig(
        platform=config_data.get('platform'),
        keyword=config_data.get('keyword'),
        max_items=config_data.get('max_items', 100),
        analysis_types=config_data.get('analysis_types', ['sentiment', 'keywords']),
        extra=config_data.get('extra', {})
    )
    
    # 创建任务
    task = scheduler.create_task(
        name=name,
        task_type=task_type_enum,
        config=task_config,
        schedule=schedule,
        interval_seconds=interval_seconds,
        enabled=True
    )
    
    return jsonify({
        'success': True,
        'message': 'Task created',
        'task': {
            'id': task.id,
            'name': task.name,
            'type': task.task_type.value,
            'status': task.status.value
        }
    }), 201


@app.route('/api/tasks/<task_id>', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_task(task_id):
    """获取任务详情"""
    task = scheduler.get_task(task_id)
    if not task:
        raise APIError('Task not found', 404)
    
    # 获取执行历史
    history = scheduler.get_task_history(task_id)
    history_data = []
    for result in history:
        history_data.append({
            'status': result.status.value,
            'started_at': result.started_at.isoformat() if result.started_at else None,
            'finished_at': result.finished_at.isoformat() if result.finished_at else None,
            'duration_seconds': result.duration_seconds,
            'items_processed': result.items_processed,
            'error': result.error_message
        })
    
    return jsonify({
        'success': True,
        'task': {
            'id': task.id,
            'name': task.name,
            'type': task.task_type.value,
            'status': task.status.value,
            'enabled': task.enabled,
            'config': {
                'platform': task.config.platform,
                'keyword': task.config.keyword,
                'max_items': task.config.max_items
            },
            'schedule': task.schedule,
            'interval_seconds': task.interval_seconds,
            'created_at': task.created_at.isoformat(),
            'last_run': task.last_run.isoformat() if task.last_run else None,
            'next_run': task.next_run.isoformat() if task.next_run else None,
            'run_count': task.run_count
        },
        'history': history_data
    })


@app.route('/api/tasks/<task_id>/run', methods=['POST'])
@require_api_key
@rate_limit(limit=10)
def run_task(task_id):
    """立即执行任务"""
    result = scheduler.run_now(task_id)
    if not result:
        raise APIError('Task not found', 404)
    
    return jsonify({
        'success': True,
        'message': 'Task executed',
        'result': {
            'status': result.status.value,
            'items_processed': result.items_processed,
            'duration_seconds': result.duration_seconds,
            'error': result.error_message
        }
    })


@app.route('/api/tasks/<task_id>/pause', methods=['POST'])
@require_api_key
@rate_limit(limit=30)
def pause_task(task_id):
    """暂停任务"""
    if scheduler.pause_task(task_id):
        return jsonify({'success': True, 'message': 'Task paused'})
    raise APIError('Task not found', 404)


@app.route('/api/tasks/<task_id>/resume', methods=['POST'])
@require_api_key
@rate_limit(limit=30)
def resume_task(task_id):
    """恢复任务"""
    if scheduler.resume_task(task_id):
        return jsonify({'success': True, 'message': 'Task resumed'})
    raise APIError('Task not found', 404)


@app.route('/api/tasks/<task_id>', methods=['DELETE'])
@require_api_key
@rate_limit(limit=30)
def delete_task(task_id):
    """删除任务"""
    if scheduler.delete_task(task_id):
        return jsonify({'success': True, 'message': 'Task deleted'})
    raise APIError('Task not found', 404)


# ============ 报告 ============

@app.route('/api/report', methods=['GET'])
@rate_limit(limit=30)
def report():
    """生成系统报告"""
    report_data = {
        'report_time': datetime.now().isoformat(),
        'system_status': 'healthy',
        'available_platforms': CollectorFactory.available_platforms(),
        'features': {
            'collection': ['微博', '知乎', '抖音', '百度', '微信公众号', '小红书', 'B站', '贴吧', '头条'],
            'processing': ['中文分词', '情感分析', '关键词提取', '命名实体识别'],
            'analysis': ['情感统计', '关系网络', '趋势分析', '风险评估']
        },
        'statistics': {
            'uptime_seconds': (datetime.now() - START_TIME).total_seconds(),
            'api_version': '1.0.0',
            'collections_count': len(collected_data_store),
            'analysis_count': len(analysis_results_store),
            'tasks_count': len(scheduler.get_all_tasks())
        }
    }
    
    return jsonify({
        'success': True,
        'report': report_data
    }), 200


@app.route('/api/dashboard/pipeline', methods=['POST'])
@require_json
@optional_api_key
@rate_limit(limit=30)
def dashboard_pipeline():
    """关键词驱动的一体化舆情流水线：采集 -> 分析 -> 词云 -> 报告。"""
    req_data = request.json or {}

    keyword = str(req_data.get('keyword', '')).strip()
    if not keyword:
        raise APIError('Missing required field: keyword', 400)

    max_items = req_data.get('max_items', 60)
    try:
        max_items = int(max_items)
    except Exception:
        max_items = 60
    max_items = max(5, min(max_items, 500))

    platforms = _normalize_platforms(req_data.get('platforms'))
    result = _run_dashboard_pipeline_internal(keyword, platforms, max_items)
    return jsonify(result), (200 if result.get('success') else 502)


@app.route('/api/reports/<pipeline_id>/export', methods=['GET'])
@optional_api_key
@rate_limit(limit=20)
def export_pipeline_report(pipeline_id):
    """导出分析报告，支持 html/docx/pdf。"""
    export_format = (request.args.get('format') or 'html').strip().lower()
    
    # 验证format
    if export_format not in ('html', 'docx', 'pdf'):
        return jsonify({'success': False, 'error': 'Invalid format. Supported: html, docx, pdf'}), 400
    
    output_path = _export_report_file(pipeline_id, export_format)
    
    if export_format == 'html':
        mimetype = 'text/html; charset=utf-8'
    elif export_format == 'docx':
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:  # pdf
        mimetype = 'application/pdf'
    
    return send_file(output_path, as_attachment=True, download_name=output_path.name, mimetype=mimetype)


@app.route('/api/monitors', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def list_monitors():
    group_id = request.args.get('group_id')
    monitors = list(monitor_profiles_store.values())
    if group_id:
        monitors = [profile for profile in monitors if profile.group_id == group_id]
    return jsonify({
        'success': True,
        'monitors': [profile.to_dict() for profile in monitors],
        'count': len(monitors)
    })


@app.route('/api/monitor-groups', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def list_monitor_groups():
    return jsonify({
        'success': True,
        'groups': [group.to_dict() for group in monitor_groups_store.values()],
        'count': len(monitor_groups_store)
    })


@app.route('/api/monitor-groups', methods=['POST'])
@require_json
@optional_api_key
@rate_limit(limit=20)
def create_monitor_group():
    data = request.json or {}
    name = str(data.get('name', '')).strip()
    if not name:
        raise APIError('Missing required field: name', 400)
    group = MonitorGroup(
        group_id=str(uuid.uuid4()),
        name=name,
        description=str(data.get('description', '') or '').strip(),
        color=str(data.get('color', '#2e8cff') or '#2e8cff').strip()
    )
    monitor_groups_store[group.group_id] = group
    return jsonify({'success': True, 'group': group.to_dict()}), 201


@app.route('/api/monitor-groups/<group_id>', methods=['PUT'])
@require_json
@optional_api_key
@rate_limit(limit=20)
def update_monitor_group(group_id):
    group = monitor_groups_store.get(group_id)
    if not group:
        raise APIError('Monitor group not found', 404)
    data = request.json or {}
    if 'name' in data:
        group.name = str(data.get('name') or '').strip() or group.name
    if 'description' in data:
        group.description = str(data.get('description') or '').strip()
    if 'color' in data:
        group.color = str(data.get('color') or group.color).strip() or group.color
    group.updated_at = datetime.now().isoformat()
    return jsonify({'success': True, 'group': group.to_dict()})


@app.route('/api/monitor-groups/<group_id>', methods=['DELETE'])
@optional_api_key
@rate_limit(limit=20)
def delete_monitor_group(group_id):
    group = monitor_groups_store.pop(group_id, None)
    if not group:
        raise APIError('Monitor group not found', 404)
    for profile in monitor_profiles_store.values():
        if profile.group_id == group_id:
            profile.group_id = None
            profile.updated_at = datetime.now().isoformat()
    return jsonify({'success': True, 'message': 'Monitor group deleted'})


@app.route('/api/monitors', methods=['POST'])
@require_json
@optional_api_key
@rate_limit(limit=20)
def create_monitor():
    data = request.json or {}
    name = str(data.get('name', '')).strip()
    keywords = _normalize_keywords(data.get('keywords'))
    platforms = _normalize_platforms(data.get('platforms'))
    if not name:
        raise APIError('Missing required field: name', 400)
    if not keywords:
        raise APIError('Missing required field: keywords', 400)

    profile = MonitorProfile(
        monitor_id=str(uuid.uuid4()),
        name=name,
        keywords=keywords,
        platforms=platforms,
        group_id=(str(data.get('group_id')).strip() if data.get('group_id') else None),
        tags=_normalize_tags(data.get('tags')),
        interval_seconds=max(60, int(data.get('interval_seconds', 1800))),
        max_items=max(5, min(int(data.get('max_items', 60)), 500)),
        thresholds=data.get('thresholds', {}),
        report_formats=data.get('report_formats', ['docx', 'pdf']),
        enabled=bool(data.get('enabled', True))
    )
    monitor_profiles_store[profile.monitor_id] = profile
    if profile.enabled:
        _schedule_monitor_profile(profile)

    return jsonify({'success': True, 'monitor': profile.to_dict()}), 201


@app.route('/api/monitors/<monitor_id>', methods=['PUT'])
@require_json
@optional_api_key
@rate_limit(limit=20)
def update_monitor(monitor_id):
    profile = monitor_profiles_store.get(monitor_id)
    if not profile:
        raise APIError('Monitor not found', 404)
    data = request.json or {}
    if 'name' in data:
        profile.name = str(data.get('name') or '').strip() or profile.name
    if 'keywords' in data:
        profile.keywords = _normalize_keywords(data.get('keywords')) or profile.keywords
    if 'platforms' in data:
        profile.platforms = _normalize_platforms(data.get('platforms'))
    if 'group_id' in data:
        profile.group_id = str(data.get('group_id') or '').strip() or None
    if 'tags' in data:
        profile.tags = _normalize_tags(data.get('tags'))
    if 'interval_seconds' in data:
        profile.interval_seconds = max(60, int(data.get('interval_seconds', profile.interval_seconds)))
    if 'max_items' in data:
        profile.max_items = max(5, min(int(data.get('max_items', profile.max_items)), 500))
    if 'thresholds' in data:
        profile.thresholds = data.get('thresholds', {})
    if 'report_formats' in data:
        profile.report_formats = data.get('report_formats') or profile.report_formats
    if 'enabled' in data:
        profile.enabled = bool(data.get('enabled'))
    profile.updated_at = datetime.now().isoformat()
    if profile.enabled:
        _schedule_monitor_profile(profile)
    else:
        _unschedule_monitor_profile(profile.monitor_id)
    return jsonify({'success': True, 'monitor': profile.to_dict()})


@app.route('/api/monitors/<monitor_id>', methods=['DELETE'])
@optional_api_key
@rate_limit(limit=20)
def delete_monitor(monitor_id):
    profile = monitor_profiles_store.pop(monitor_id, None)
    if not profile:
        raise APIError('Monitor not found', 404)
    _unschedule_monitor_profile(monitor_id)
    return jsonify({'success': True, 'message': 'Monitor deleted'})


@app.route('/api/monitors/<monitor_id>/run', methods=['POST'])
@optional_api_key
@rate_limit(limit=20)
def run_monitor(monitor_id):
    result = _run_monitor_profile_once(monitor_id)
    if not result.get('success'):
        return jsonify(result), 404
    return jsonify(result)


@app.route('/api/dashboard/education', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def education_dashboard_data():
    """教育舆情大屏聚合数据"""
    province_keywords = [
        '北京', '上海', '天津', '重庆', '河北', '山西', '辽宁', '吉林', '黑龙江', '江苏',
        '浙江', '安徽', '福建', '江西', '山东', '河南', '湖北', '湖南', '广东', '海南',
        '四川', '贵州', '云南', '陕西', '甘肃', '青海', '台湾', '内蒙古', '广西', '西藏',
        '宁夏', '新疆', '香港', '澳门'
    ]

    collection_records = sorted(
        collected_data_store.values(),
        key=lambda x: x.get('collected_at', ''),
        reverse=True
    )

    recent_items = []
    for record in collection_records[:30]:
        for item in record.get('items', [])[:30]:
            merged = {
                'platform': record.get('platform', item.get('platform', 'unknown')),
                'keyword': record.get('keyword', ''),
                **item
            }
            recent_items.append(merged)

    recent_items = recent_items[:400]

    platform_counter = Counter()
    province_counter = Counter()
    sentiment_counter = Counter({'positive': 0, 'neutral': 0, 'negative': 0})
    trend_counter = defaultdict(int)
    heat_total = 0

    for item in recent_items:
        platform = item.get('platform') or 'unknown'
        platform_counter[platform] += 1

        content = (item.get('content') or '').strip()
        for province in province_keywords:
            if province in content:
                province_counter[province] += 1

        likes = int(item.get('likes') or 0)
        comments = int(item.get('comments') or 0)
        shares = int(item.get('shares') or 0)
        heat_total += max(1, likes + comments * 2 + shares * 3)

        if content:
            try:
                result = text_processor.sentiment_analysis(content)
                sentiment_counter[result.sentiment] += 1
            except Exception:
                sentiment_counter['neutral'] += 1

        time_text = item.get('publish_time') or item.get('collected_at')
        if time_text:
            try:
                dt = datetime.fromisoformat(str(time_text).replace('Z', '+00:00'))
                hour_key = f"{dt.hour:02d}:00"
                trend_counter[hour_key] += 1
            except Exception:
                pass

    if not province_counter:
        province_counter.update({
            '北京': 26, '上海': 28, '广东': 24, '江苏': 20,
            '浙江': 18, '山东': 16, '四川': 14, '湖北': 13
        })

    if not trend_counter:
        trend_counter.update({'17:00': 6, '18:00': 8, '19:00': 7, '20:00': 9, '21:00': 8, '22:00': 7, '23:00': 7})

    top_news = []
    for item in recent_items[:12]:
        text = (item.get('content') or '').replace('\n', ' ').strip()
        if not text:
            continue
        top_news.append(text[:80])

    if not top_news:
        top_news = [
            '暂未采集到实时舆情数据，请先创建采集任务。',
            '建议监控关键词：师资、收费、校园安全、课程质量。'
        ]

    active_alerts = []
    try:
        active_alerts = alert_manager.get_active_alerts()
    except Exception:
        active_alerts = []

    latest_pipeline = None
    if dashboard_pipeline_store:
        latest_pipeline = max(
            dashboard_pipeline_store.values(),
            key=lambda x: x.get('created_at', '')
        )

    latest_wordcloud = latest_pipeline.get('wordcloud', []) if latest_pipeline else []
    latest_report = latest_pipeline.get('report', {}) if latest_pipeline else {}

    alerts_payload = []
    for alert in (active_alerts or [])[:5]:
        if isinstance(alert, dict):
            alerts_payload.append(alert)
        elif hasattr(alert, 'to_dict'):
            alerts_payload.append(alert.to_dict())

    return jsonify({
        'success': True,
        'dashboard': {
            'updated_at': datetime.now().isoformat(),
            'overview': {
                'total_heat': heat_total,
                'collections_count': len(collected_data_store),
                'analysis_count': len(analysis_results_store),
                'active_alerts': len(active_alerts or [])
            },
            'news': top_news,
            'platform_distribution': dict(platform_counter),
            'sentiment': {
                'positive': sentiment_counter['positive'],
                'neutral': sentiment_counter['neutral'],
                'negative': sentiment_counter['negative']
            },
            'trend': dict(sorted(trend_counter.items(), key=lambda x: x[0])),
            'province_heat': [
                {'name': name, 'value': value * 120}
                for name, value in province_counter.most_common()
            ],
            'alerts': alerts_payload,
            'wordcloud': latest_wordcloud[:40],
            'latest_report': latest_report
        }
    })


@app.route('/api/dashboard/groups-overview', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def dashboard_groups_overview():
    """监控分组总览数据。"""
    return jsonify(_build_group_overview_payload())


@app.route('/api/dashboard/competitor-overview', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def dashboard_competitor_overview():
    """监控分组竞品对比数据。"""
    base_group_id = request.args.get('base_group_id', '').strip() or None
    return jsonify(_build_competitor_overview_payload(base_group_id=base_group_id))


# ============ 分析结果查询 ============

@app.route('/api/analysis/<analysis_id>', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_analysis(analysis_id):
    """获取分析结果"""
    result = analysis_results_store.get(analysis_id)
    if not result:
        raise APIError('Analysis not found', 404)
    
    return jsonify({
        'success': True,
        'analysis_id': analysis_id,
        'created_at': result['created_at'],
        'source_count': result['source_count'],
        'results': result['results']
    })


# ============ API Key 管理（仅管理员）============

@app.route('/api/keys', methods=['POST'])
@require_api_key
@rate_limit(limit=10)
def create_api_key():
    """创建新的 API Key"""
    from osint_cn.security import api_key_manager
    
    # 检查是否是管理员
    if g.api_key_info.get('role') != 'admin':
        raise APIError('Admin role required', 403)
    
    data = request.json or {}
    name = data.get('name', 'new_key')
    role = data.get('role', 'user')
    rate_limit_val = data.get('rate_limit', 100)
    
    new_key = api_key_manager.generate_key(name, role, rate_limit_val)
    
    return jsonify({
        'success': True,
        'message': 'API Key created',
        'api_key': new_key,
        'note': '请妥善保存此 API Key，它不会再次显示'
    }), 201


# ============ 实时采集 API ============

@app.route('/api/realtime/tasks', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def list_realtime_tasks():
    """列出实时采集任务"""
    tasks = realtime_collector.list_tasks()
    return jsonify({
        'success': True,
        'tasks': tasks,
        'stats': realtime_collector.get_stats()
    })


@app.route('/api/realtime/tasks', methods=['POST'])
@require_api_key
@require_json
@rate_limit(limit=20)
def create_realtime_task():
    """创建实时采集任务"""
    data = request.json
    
    platforms = data.get('platforms', ['weibo'])
    keywords = data.get('keywords', [])
    interval = data.get('interval_seconds', 60)
    max_items = data.get('max_items', 100)
    
    if not keywords:
        raise APIError('keywords 不能为空', 400)
    
    task_id = realtime_collector.create_task(
        platforms=platforms,
        keywords=keywords,
        interval_seconds=interval,
        max_items=max_items
    )
    
    return jsonify({
        'success': True,
        'task_id': task_id,
        'message': '实时采集任务已创建'
    }), 201


@app.route('/api/realtime/tasks/<task_id>', methods=['DELETE'])
@require_api_key
@rate_limit(limit=20)
def delete_realtime_task(task_id):
    """删除实时采集任务"""
    if realtime_collector.remove_task(task_id):
        return jsonify({'success': True, 'message': '任务已删除'})
    raise APIError('任务不存在', 404)


@app.route('/api/realtime/tasks/<task_id>/pause', methods=['POST'])
@require_api_key
@rate_limit(limit=20)
def pause_realtime_task(task_id):
    """暂停实时采集任务"""
    if realtime_collector.pause_task(task_id):
        return jsonify({'success': True, 'message': '任务已暂停'})
    raise APIError('任务不存在', 404)


@app.route('/api/realtime/tasks/<task_id>/resume', methods=['POST'])
@require_api_key
@rate_limit(limit=20)
def resume_realtime_task(task_id):
    """恢复实时采集任务"""
    if realtime_collector.resume_task(task_id):
        return jsonify({'success': True, 'message': '任务已恢复'})
    raise APIError('任务不存在', 404)


@app.route('/api/realtime/data', methods=['GET'])
@optional_api_key
@rate_limit(limit=100)
def get_realtime_data():
    """获取实时采集数据"""
    limit = request.args.get('limit', 100, type=int)
    data = realtime_collector.get_buffer_data(max_items=limit)
    return jsonify({
        'success': True,
        'data': data,
        'count': len(data)
    })


# ============ 情报分析 API ============

@app.route('/api/intel/analyze', methods=['POST'])
@require_api_key
@require_json
@rate_limit(limit=60)
def analyze_intelligence():
    """分析情报"""
    data = request.json
    content = data.get('content', '')
    source = data.get('source', 'unknown')
    source_url = data.get('source_url', '')
    intel_type = data.get('intel_type', 'osint')
    source_category = data.get('source_category', 'unknown')
    
    if not content:
        raise APIError('content 不能为空', 400)
    
    # 转换情报类型
    try:
        intel_type_enum = IntelligenceType(intel_type)
    except:
        intel_type_enum = IntelligenceType.OSINT
    
    result = intelligence_analyzer.analyze(
        content=content,
        source=source,
        source_url=source_url,
        intel_type=intel_type_enum,
        source_category=source_category,
        metadata=data.get('metadata')
    )
    
    return jsonify({
        'success': True,
        'intel': result.to_dict()
    })


@app.route('/api/intel/search', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def search_intelligence():
    """搜索情报"""
    keyword = request.args.get('keyword')
    category = request.args.get('category')
    threat_level = request.args.get('threat_level')
    limit = request.args.get('limit', 50, type=int)
    
    # 转换参数
    category_enum = None
    if category:
        try:
            category_enum = IntelligenceCategory(category)
        except:
            pass
    
    threat_level_enum = None
    if threat_level:
        try:
            threat_level_enum = ThreatLevel[threat_level.upper()]
        except:
            pass
    
    results = intelligence_analyzer.search_intel(
        keyword=keyword,
        category=category_enum,
        threat_level=threat_level_enum,
        limit=limit
    )
    
    return jsonify({
        'success': True,
        'results': results,
        'count': len(results)
    })


@app.route('/api/intel/<intel_id>', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_intelligence_detail(intel_id):
    """获取情报详情"""
    detail = intelligence_analyzer.get_intel(intel_id)
    if not detail:
        raise APIError('情报不存在', 404)
    return jsonify({
        'success': True,
        'intel': detail
    })


@app.route('/api/intel/report', methods=['POST'])
@require_api_key
@require_json
@rate_limit(limit=10)
def generate_intel_report():
    """生成态势报告"""
    data = request.json
    period_hours = data.get('period_hours', 24)
    title = data.get('title')
    
    report = intelligence_analyzer.generate_situation_report(
        period_hours=period_hours,
        title=title
    )
    
    return jsonify({
        'success': True,
        'report': report.to_dict()
    })


@app.route('/api/intel/dashboard', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_intel_dashboard():
    """获取情报仪表板数据"""
    data = intelligence_analyzer.get_dashboard_data()
    return jsonify({
        'success': True,
        'data': data
    })


@app.route('/api/intel/stats', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_intel_stats():
    """获取情报统计"""
    stats = intelligence_analyzer.get_stats()
    return jsonify({
        'success': True,
        'stats': stats
    })


# ============ 威胁指标 (IOC) API ============

@app.route('/api/ioc/add', methods=['POST'])
@require_api_key
@require_json
@rate_limit(limit=100)
def add_threat_indicator():
    """添加威胁指标"""
    data = request.json
    
    indicator_type = data.get('type')
    value = data.get('value')
    
    if not indicator_type or not value:
        raise APIError('type 和 value 必填', 400)
    
    indicator_id = intelligence_analyzer.threat_intel.add_indicator(
        indicator_type=indicator_type,
        value=value,
        threat_type=data.get('threat_type', 'unknown'),
        confidence=data.get('confidence', 0.5),
        source=data.get('source', 'api'),
        tags=data.get('tags', [])
    )
    
    return jsonify({
        'success': True,
        'indicator_id': indicator_id
    }), 201


@app.route('/api/ioc/check', methods=['POST'])
@optional_api_key
@require_json
@rate_limit(limit=200)
def check_threat_indicator():
    """检查威胁指标"""
    data = request.json
    
    indicator_type = data.get('type')
    value = data.get('value')
    
    if not indicator_type or not value:
        raise APIError('type 和 value 必填', 400)
    
    indicator = intelligence_analyzer.threat_intel.check_indicator(indicator_type, value)
    
    return jsonify({
        'success': True,
        'found': indicator is not None,
        'indicator': indicator.to_dict() if indicator else None
    })


@app.route('/api/ioc/search', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def search_threat_indicators():
    """搜索威胁指标"""
    indicator_type = request.args.get('type')
    value_pattern = request.args.get('pattern')
    min_confidence = request.args.get('min_confidence', 0.0, type=float)
    
    results = intelligence_analyzer.threat_intel.search_indicators(
        indicator_type=indicator_type,
        value_pattern=value_pattern,
        min_confidence=min_confidence
    )
    
    return jsonify({
        'success': True,
        'indicators': results,
        'count': len(results)
    })


@app.route('/api/ioc/stats', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_ioc_stats():
    """获取威胁指标统计"""
    stats = intelligence_analyzer.threat_intel.get_stats()
    return jsonify({
        'success': True,
        'stats': stats
    })


# ============ 关系挖掘 API ============

@app.route('/api/graph/add', methods=['POST'])
@require_api_key
@require_json
@rate_limit(limit=60)
def add_to_graph():
    """添加文档到知识图谱"""
    data = request.json
    text = data.get('text', '')
    source = data.get('source', 'unknown')
    author = data.get('author')
    
    if not text:
        raise APIError('text 不能为空', 400)
    
    result = knowledge_graph.add_document(
        text=text,
        source=source,
        author=author,
        metadata=data.get('metadata')
    )
    
    return jsonify({
        'success': True,
        'result': result
    })


@app.route('/api/graph/entity/<name>', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def query_entity(name):
    """查询实体"""
    result = knowledge_graph.query_entity(name)
    if not result:
        raise APIError('实体不存在', 404)
    return jsonify({
        'success': True,
        'result': result
    })


@app.route('/api/graph/subgraph/<entity_id>', methods=['GET'])
@optional_api_key
@rate_limit(limit=30)
def get_subgraph(entity_id):
    """获取子图"""
    depth = request.args.get('depth', 2, type=int)
    max_nodes = request.args.get('max_nodes', 100, type=int)
    
    result = knowledge_graph.get_subgraph(entity_id, depth=depth, max_nodes=max_nodes)
    return jsonify({
        'success': True,
        'subgraph': result
    })


@app.route('/api/graph/path', methods=['GET'])
@optional_api_key
@rate_limit(limit=30)
def find_entity_path():
    """查找实体之间的路径"""
    source_id = request.args.get('source')
    target_id = request.args.get('target')
    
    if not source_id or not target_id:
        raise APIError('source 和 target 参数必填', 400)
    
    path = knowledge_graph.find_path(source_id, target_id)
    return jsonify({
        'success': True,
        'path': path,
        'found': path is not None
    })


@app.route('/api/graph/stats', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_graph_stats():
    """获取图谱统计"""
    stats = knowledge_graph.get_stats()
    return jsonify({
        'success': True,
        'stats': stats
    })


@app.route('/api/graph/export/neo4j', methods=['GET'])
@require_api_key
@rate_limit(limit=10)
def export_to_neo4j():
    """导出为 Neo4j 格式"""
    data = knowledge_graph.export_for_neo4j()
    return jsonify({
        'success': True,
        'data': data
    })


# ============ 社交网络分析 API ============

@app.route('/api/social/user', methods=['POST'])
@require_api_key
@require_json
@rate_limit(limit=60)
def add_social_user():
    """添加社交用户"""
    data = request.json
    
    user_id = data.get('user_id')
    username = data.get('username')
    platform = data.get('platform', 'unknown')
    
    if not user_id or not username:
        raise APIError('user_id 和 username 必填', 400)
    
    node = social_network.add_user(
        user_id=user_id,
        username=username,
        platform=platform,
        followers=data.get('followers', 0),
        following=data.get('following', 0),
        posts=data.get('posts', 0)
    )
    
    return jsonify({
        'success': True,
        'user': node.to_dict()
    })


@app.route('/api/social/interaction', methods=['POST'])
@require_api_key
@require_json
@rate_limit(limit=100)
def add_social_interaction():
    """添加社交互动"""
    data = request.json
    
    source_id = data.get('source_id')
    target_id = data.get('target_id')
    weight = data.get('weight', 1.0)
    
    if not source_id or not target_id:
        raise APIError('source_id 和 target_id 必填', 400)
    
    social_network.add_interaction(source_id, target_id, weight)
    
    return jsonify({
        'success': True,
        'message': '互动关系已添加'
    })


@app.route('/api/social/influence', methods=['GET'])
@optional_api_key
@rate_limit(limit=30)
def calculate_influence():
    """计算影响力"""
    ranks = social_network.calculate_pagerank()
    return jsonify({
        'success': True,
        'ranks': {k: round(v, 6) for k, v in sorted(ranks.items(), key=lambda x: x[1], reverse=True)[:50]}
    })


@app.route('/api/social/communities', methods=['GET'])
@optional_api_key
@rate_limit(limit=30)
def detect_communities():
    """社区检测"""
    min_size = request.args.get('min_size', 3, type=int)
    communities = social_network.detect_communities(min_community_size=min_size)
    return jsonify({
        'success': True,
        'communities': communities,
        'count': len(communities)
    })


@app.route('/api/social/connectors', methods=['GET'])
@optional_api_key
@rate_limit(limit=30)
def find_key_connectors():
    """找出关键连接者"""
    top_n = request.args.get('top', 10, type=int)
    connectors = social_network.find_key_connectors(top_n=top_n)
    return jsonify({
        'success': True,
        'connectors': connectors
    })


@app.route('/api/social/network/<user_id>', methods=['GET'])
@optional_api_key
@rate_limit(limit=30)
def get_user_network(user_id):
    """获取用户社交网络"""
    depth = request.args.get('depth', 2, type=int)
    network = social_network.get_user_network(user_id, depth=depth)
    return jsonify({
        'success': True,
        'network': network
    })


@app.route('/api/social/stats', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_social_stats():
    """获取社交网络统计"""
    stats = social_network.get_stats()
    return jsonify({
        'success': True,
        'stats': stats
    })


# ============ 风险预警 API ============

@app.route('/api/alert/rules', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def list_alert_rules():
    """列出预警规则"""
    rules = alert_manager.rule_engine.list_rules()
    return jsonify({
        'success': True,
        'rules': rules,
        'count': len(rules)
    })


@app.route('/api/alert/rules', methods=['POST'])
@require_api_key
@require_json
@rate_limit(limit=20)
def create_alert_rule():
    """创建预警规则"""
    data = request.json
    
    name = data.get('name')
    rule_type = data.get('type', 'keyword')
    level = data.get('level', 'WARNING')
    
    if not name:
        raise APIError('name 必填', 400)
    
    config = {
        'keywords': data.get('keywords', []),
        'regex_pattern': data.get('regex_pattern'),
        'threshold_field': data.get('threshold_field'),
        'threshold_value': data.get('threshold_value', 0),
        'threshold_operator': data.get('threshold_operator', '>'),
        'sentiment_threshold': data.get('sentiment_threshold', -0.5),
        'frequency_count': data.get('frequency_count', 10),
        'frequency_window_minutes': data.get('frequency_window_minutes', 60),
        'notify_channels': data.get('notify_channels', []),
        'cooldown_minutes': data.get('cooldown_minutes', 30),
        'description': data.get('description', '')
    }
    
    rule_id = alert_manager.create_custom_rule(name, rule_type, level, config)
    
    return jsonify({
        'success': True,
        'rule_id': rule_id,
        'message': '预警规则已创建'
    }), 201


@app.route('/api/alert/rules/<rule_id>', methods=['DELETE'])
@require_api_key
@rate_limit(limit=20)
def delete_alert_rule(rule_id):
    """删除预警规则"""
    if alert_manager.rule_engine.remove_rule(rule_id):
        return jsonify({'success': True, 'message': '规则已删除'})
    raise APIError('规则不存在', 404)


@app.route('/api/alert/rules/<rule_id>/enable', methods=['POST'])
@require_api_key
@rate_limit(limit=20)
def enable_alert_rule(rule_id):
    """启用预警规则"""
    if alert_manager.rule_engine.enable_rule(rule_id):
        return jsonify({'success': True, 'message': '规则已启用'})
    raise APIError('规则不存在', 404)


@app.route('/api/alert/rules/<rule_id>/disable', methods=['POST'])
@require_api_key
@rate_limit(limit=20)
def disable_alert_rule(rule_id):
    """禁用预警规则"""
    if alert_manager.rule_engine.disable_rule(rule_id):
        return jsonify({'success': True, 'message': '规则已禁用'})
    raise APIError('规则不存在', 404)


@app.route('/api/alert/check', methods=['POST'])
@require_api_key
@require_json
@rate_limit(limit=100)
def check_content_alert():
    """检查内容是否触发预警"""
    data = request.json
    content = data.get('content', '')
    source = data.get('source', 'api')
    
    if not content:
        raise APIError('content 不能为空', 400)
    
    alerts = alert_manager.process_content(
        content=content,
        source=source,
        metadata=data.get('metadata')
    )
    
    return jsonify({
        'success': True,
        'triggered': len(alerts) > 0,
        'alerts': [a.to_dict() for a in alerts]
    })


@app.route('/api/alert/active', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_active_alerts():
    """获取活跃预警"""
    level = request.args.get('level')
    level_enum = AlertLevel[level.upper()] if level else None
    
    alerts = alert_manager.get_active_alerts(level=level_enum)
    return jsonify({
        'success': True,
        'alerts': alerts,
        'count': len(alerts)
    })


@app.route('/api/alert/history', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_alert_history():
    """获取预警历史"""
    limit = request.args.get('limit', 100, type=int)
    level = request.args.get('level')
    level_enum = AlertLevel[level.upper()] if level else None
    
    history = alert_manager.get_alert_history(limit=limit, level=level_enum)
    return jsonify({
        'success': True,
        'alerts': history,
        'count': len(history)
    })


@app.route('/api/alert/<alert_id>/acknowledge', methods=['POST'])
@require_api_key
@rate_limit(limit=60)
def acknowledge_alert(alert_id):
    """确认预警"""
    user = request.json.get('user', 'api_user') if request.json else 'api_user'
    
    if alert_manager.acknowledge_alert(alert_id, user):
        return jsonify({'success': True, 'message': '预警已确认'})
    raise APIError('预警不存在', 404)


@app.route('/api/alert/<alert_id>/resolve', methods=['POST'])
@require_api_key
@rate_limit(limit=60)
def resolve_alert(alert_id):
    """解决预警"""
    if alert_manager.resolve_alert(alert_id):
        return jsonify({'success': True, 'message': '预警已解决'})
    raise APIError('预警不存在', 404)


@app.route('/api/alert/stats', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_alert_stats():
    """获取预警统计"""
    stats = alert_manager.get_stats()
    return jsonify({
        'success': True,
        'stats': stats
    })


# ============ 综合情报接口 ============

@app.route('/api/intelligence/process', methods=['POST'])
@require_api_key
@require_json
@rate_limit(limit=30)
def process_intelligence():
    """综合情报处理 - 一站式分析"""
    data = request.json
    content = data.get('content', '')
    source = data.get('source', 'unknown')
    source_url = data.get('source_url', '')
    author = data.get('author')
    
    if not content:
        raise APIError('content 不能为空', 400)
    
    results = {
        'content_id': str(uuid.uuid4())[:8],
        'processed_at': datetime.now().isoformat()
    }
    
    # 1. 情报分析
    intel_result = intelligence_analyzer.analyze(
        content=content,
        source=source,
        source_url=source_url,
        metadata={'author': author}
    )
    results['intelligence'] = intel_result.to_dict()
    
    # 2. 实体识别和关系抽取
    results['graph'] = knowledge_graph.add_document(
        text=content,
        source=source,
        author=author
    )
    
    # 3. 风险预警检查
    alerts = alert_manager.process_content(
        content=content,
        source=source,
        metadata={'author': author}
    )
    results['alerts'] = {
        'triggered': len(alerts) > 0,
        'count': len(alerts),
        'alerts': [a.to_dict() for a in alerts]
    }
    
    return jsonify({
        'success': True,
        'results': results
    })


@app.route('/api/intelligence/dashboard', methods=['GET'])
@optional_api_key
@rate_limit(limit=30)
def get_intelligence_dashboard():
    """情报仪表板 - 综合概览"""
    dashboard = {
        'updated_at': datetime.now().isoformat(),
        'realtime': realtime_collector.get_stats(),
        'intelligence': intelligence_analyzer.get_dashboard_data(),
        'graph': knowledge_graph.get_stats(),
        'social': social_network.get_stats(),
        'alerts': alert_manager.get_stats()
    }
    
    return jsonify({
        'success': True,
        'dashboard': dashboard
    })


# =============================================================================
# 批量采集任务 API
# =============================================================================

@app.route('/api/batch/collect', methods=['POST'])
@optional_api_key
@rate_limit(limit=20)
def batch_collect():
    """
    批量创建采集任务
    
    请求体:
    {
        "name": "教育行业舆情监测",
        "tasks": [
            {"platform": "weibo", "keyword": "在线教育", "max_items": 20},
            {"platform": "zhihu", "keyword": "教育改革", "max_items": 15},
            {"platform": "douyin", "keyword": "职业培训", "max_items": 10}
        ],
        "auto_start": true
    }
    
    响应:
    {
        "success": true,
        "task_id": "uuid",
        "name": "教育行业舆情监测",
        "subtasks_count": 3,
        "status": "running",
        "created_at": "2026-03-08T15:00:00"
    }
    """
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'success': False,
                'error': '请求体为空'
            }), 400
        
        name = data.get('name', f'批量任务_{datetime.now().strftime("%Y%m%d_%H%M%S")}')
        tasks = data.get('tasks', [])
        auto_start = data.get('auto_start', True)
        metadata = data.get('metadata', {})
        
        if not tasks:
            return jsonify({
                'success': False,
                'error': 'tasks 不能为空'
            }), 400
        
        # 验证每个子任务
        for i, task in enumerate(tasks):
            if 'platform' not in task:
                return jsonify({
                    'success': False,
                    'error': f'第 {i+1} 个任务缺少 platform 字段'
                }), 400
            
            if 'keyword' not in task:
                return jsonify({
                    'success': False,
                    'error': f'第 {i+1} 个任务缺少 keyword 字段'
                }), 400
            
            # 验证平台是否支持
            try:
                Platform(task['platform'])
            except ValueError:
                return jsonify({
                    'success': False,
                    'error': f'不支持的平台: {task["platform"]}'
                }), 400
        
        # 获取调度器
        scheduler = get_scheduler()
        
        # 创建批量任务
        batch_task = scheduler.create_batch_task(
            name=name,
            platforms_keywords=tasks,
            metadata=metadata
        )
        
        # 自动启动
        if auto_start:
            scheduler.submit_task(batch_task.id)
        
        return jsonify({
            'success': True,
            'task_id': batch_task.id,
            'name': batch_task.name,
            'subtasks_count': len(batch_task.subtasks),
            'status': batch_task.status.value,
            'created_at': batch_task.created_at.isoformat()
        })
    
    except Exception as e:
        logger.error(f"创建批量任务失败: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/batch/tasks', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def list_batch_tasks():
    """
    获取批量任务列表
    
    查询参数:
    - status: 按状态过滤 (pending/running/completed/failed/cancelled)
    - limit: 返回数量 (默认50)
    
    响应:
    {
        "success": true,
        "tasks": [
            {
                "id": "uuid",
                "name": "教育行业舆情监测",
                "status": "completed",
                "created_at": "2026-03-08T15:00:00",
                "completed_at": "2026-03-08T15:05:23",
                "total_items": 45,
                "success_count": 3,
                "failed_count": 0,
                "subtasks": [...]
            }
        ],
        "count": 1
    }
    """
    try:
        status_str = request.args.get('status')
        limit = int(request.args.get('limit', 50))
        
        status_filter = None
        if status_str:
            try:
                status_filter = BatchTaskStatus(status_str)
            except ValueError:
                return jsonify({
                    'success': False,
                    'error': f'无效的状态值: {status_str}'
                }), 400
        
        # 获取调度器
        scheduler = get_scheduler()
        
        # 获取任务列表
        tasks = scheduler.list_tasks(status=status_filter, limit=limit)
        
        return jsonify({
            'success': True,
            'tasks': [task.to_dict() for task in tasks],
            'count': len(tasks)
        })
    
    except Exception as e:
        logger.error(f"获取批量任务列表失败: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/batch/tasks/<task_id>', methods=['GET'])
@optional_api_key
@rate_limit(limit=60)
def get_batch_task(task_id):
    """
    获取批量任务详情
    
    响应:
    {
        "success": true,
        "task": {
            "id": "uuid",
            "name": "教育行业舆情监测",
            "status": "completed",
            "created_at": "2026-03-08T15:00:00",
            "started_at": "2026-03-08T15:00:01",
            "completed_at": "2026-03-08T15:05:23",
            "total_items": 45,
            "success_count": 3,
            "failed_count": 0,
            "subtasks": [
                {
                    "id": "subtask-uuid",
                    "platform": "weibo",
                    "keyword": "在线教育",
                    "max_items": 20,
                    "status": "completed",
                    "items_collected": 18,
                    "collection_id": "collection-uuid",
                    "started_at": "2026-03-08T15:00:01",
                    "completed_at": "2026-03-08T15:01:45"
                }
            ]
        }
    }
    """
    try:
        # 获取调度器
        scheduler = get_scheduler()
        
        # 获取任务
        task = scheduler.get_task(task_id)
        
        if not task:
            return jsonify({
                'success': False,
                'error': '任务不存在'
            }), 404
        
        return jsonify({
            'success': True,
            'task': task.to_dict()
        })
    
    except Exception as e:
        logger.error(f"获取批量任务详情失败: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/batch/tasks/<task_id>/start', methods=['POST'])
@optional_api_key
@rate_limit(limit=20)
def start_batch_task(task_id):
    """
    启动批量任务
    
    响应:
    {
        "success": true,
        "task_id": "uuid",
        "status": "running"
    }
    """
    try:
        # 获取调度器
        scheduler = get_scheduler()
        
        # 提交任务
        success = scheduler.submit_task(task_id)
        
        if not success:
            return jsonify({
                'success': False,
                'error': '任务启动失败,可能任务不存在或状态不允许启动'
            }), 400
        
        task = scheduler.get_task(task_id)
        
        return jsonify({
            'success': True,
            'task_id': task_id,
            'status': task.status.value if task else 'unknown'
        })
    
    except Exception as e:
        logger.error(f"启动批量任务失败: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/batch/tasks/<task_id>/cancel', methods=['POST'])
@optional_api_key
@rate_limit(limit=20)
def cancel_batch_task(task_id):
    """
    取消批量任务
    
    响应:
    {
        "success": true,
        "task_id": "uuid",
        "status": "cancelled"
    }
    """
    try:
        # 获取调度器
        scheduler = get_scheduler()
        
        # 取消任务
        success = scheduler.cancel_task(task_id)
        
        if not success:
            return jsonify({
                'success': False,
                'error': '任务取消失败,可能任务不存在或状态不允许取消'
            }), 400
        
        task = scheduler.get_task(task_id)
        
        return jsonify({
            'success': True,
            'task_id': task_id,
            'status': task.status.value if task else 'unknown'
        })
    
    except Exception as e:
        logger.error(f"取消批量任务失败: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/batch/stats', methods=['GET'])
@optional_api_key
@rate_limit(limit=100)
def get_batch_stats():
    """
    获取批量采集调度器统计信息
    
    响应:
    {
        "success": true,
        "stats": {
            "scheduler_running": true,
            "max_workers": 3,
            "active_workers": 3,
            "queue_size": 5,
            "total_tasks": 10,
            "tasks_by_status": {
                "pending": 2,
                "running": 3,
                "completed": 4,
                "failed": 1,
                "cancelled": 0
            },
            "total_items_collected": 256
        }
    }
    """
    try:
        # 获取调度器
        scheduler = get_scheduler()
        
        # 获取统计信息
        stats = scheduler.get_statistics()
        
        return jsonify({
            'success': True,
            'stats': stats
        })
    
    except Exception as e:
        logger.error(f"获取批量统计信息失败: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5002)